# -*- coding: utf-8 -*-
"""
docx-toolkit.pyw - docx-toolkit 统一入口（单文件版，所有功能内联）

Markdown -> Word 文档转换工具集。

子命令:
  to-docx   Markdown -> Word      (tech/paper 双样式; 两者排版一致，仅表格不同: tech=深蓝网格表 / paper=三线表)

双击运行 -> GUI 模式：
  第一行 单选：MD→Word / MD→论文版 两种（排版完全一致，仅表格样式不同）
  第二行 批量模式复选框（勾选后第一行所选功能对整个目录同类文件批量转换，与第一行复选）

说明：图片已 base64 内联进 md（data URI），转换时无需外部资源路径。
"""

import argparse
import os
import sys
import re
import shutil
import subprocess
import tempfile
import glob

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# 共享工具函数（MD→Word 双样式共用，仅一份）
# ============================================================
def find_pandoc():
    found = shutil.which("pandoc")
    if found:
        return found
    candidates = [
        os.path.join(os.path.expanduser("~"), "bin", "pandoc.exe"),
        r"C:\Users\honghuang\bin\pandoc.exe",
        r"C:\Program Files\Pandoc\pandoc.exe",
    ]
    for c in candidates:
        if c and os.path.isfile(c):
            return c
    return None


def set_cell_shading(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_border(cell, color='D1D5DB', width=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_elm = OxmlElement(f'w:{edge}')
        edge_elm.set(qn('w:val'), 'single')
        edge_elm.set(qn('w:sz'), str(width))
        edge_elm.set(qn('w:color'), color)
        tcBorders.append(edge_elm)
    tcPr.append(tcBorders)


def _apply_grid_widths(table, twips):
    """把列宽（twips 列表）同步写入 tblGrid 的 gridCol 与 tblPr 内的 tblW，
    使列宽真正生效（cell.width 只写 tcW，Word 以 gridCol 为准）。

    注意：tblW 必须位于 tblPr 内部才生效。早期实现用 tbl.find(qn('w:tblW'))
    可能找到非 tblPr 内的节点或建出孤立 tblW（直接挂在 w:tbl 下），
    导致真正的 tblPr/tblW 仍为 auto/0 → 列宽为零、单元格竖排、看似“无表格”。
    这里固定定位 tblPr 内的 tblW，并清掉任何孤立 tblW。"""
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    q = lambda t: '{%s}%s' % (W, t)
    tbl = table._tbl
    # 1) gridCol
    grid = tbl.find(q('tblGrid'))
    if grid is None:
        grid = OxmlElement(q('tblGrid'))
        tbl.append(grid)
    gridCols = grid.findall(q('gridCol'))
    for j, w in enumerate(twips):
        if j < len(gridCols):
            gridCols[j].set(q('w'), str(w))
        else:
            gc = OxmlElement(q('gridCol'))
            gc.set(q('w'), str(w))
            grid.append(gc)
    # 2) tblPr 必须存在
    pr = tbl.find(q('tblPr'))
    if pr is None:
        pr = OxmlElement(q('tblPr'))
        tbl.insert(0, pr)
    # 3) 清掉挂在 w:tbl 下（非 tblPr 内）的孤立 tblW
    for stray in tbl.findall(q('tblW')):
        tbl.remove(stray)
    # 4) 定位 / 新建 tblPr 内的 tblW
    tblW = pr.find(q('tblW'))
    if tblW is None:
        tblW = OxmlElement(q('tblW'))
        pr.append(tblW)
    tblW.set(q('type'), 'dxa')
    tblW.set(q('w'), str(sum(twips)))


def _is_equation_table(table):
    """判断是否为公式编号表格（2列+含 oMathPara/oMath OMML）。"""
    if len(table.columns) != 2:
        return False
    tbl_xml = table._tbl
    return (tbl_xml.find('.//%s' % _qn_m('oMathPara')) is not None or
            tbl_xml.find('.//%s' % _qn_m('oMath')) is not None)


# ---- 表格列宽：固定撑满版心（≈16cm），各列等宽分布（不使用内容自适应）----
def _compute_autofit_widths(table, text_w_twips):
    """表格固定撑满版心宽度 text_w_twips（≈16cm），各列按等宽分布。
    不做内容自适应，恢复为简单固定的全宽表格，避免表格宽度反复调整。"""
    n_cols = len(table.columns)
    if n_cols == 0:
        return [text_w_twips]
    base = text_w_twips // n_cols
    widths = [base] * n_cols
    widths[-1] += text_w_twips - base * n_cols   # 余数补到末列，保证总和精确
    return widths


def style_tables(doc):
    """深蓝表头 + 灰边 表格样式（MD→Word 技术文档用，论文版改用三线表）。
    跳过公式编号表（由 number_equations 单独处理，需 auto 列宽避免折行）。"""
    for table in doc.tables:
        if _is_equation_table(table):
            continue
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        try:
            table.autofit = False
        except Exception:
            pass

        hdr_cells = table.rows[0].cells
        for j, cell in enumerate(hdr_cells):
            set_cell_shading(cell, '2E5A9C')
            set_cell_border(cell, color='2E5A9C', width=8)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.bold = True
                    r.font.size = Pt(9)              # 小五，表头与数据统一 9pt
                    r.font.color.rgb = RGBColor(255, 255, 255)
                    r.font.name = '黑体'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        for i in range(1, len(table.rows)):
            row_cells = table.rows[i].cells
            for j, cell in enumerate(row_cells):
                set_cell_border(cell, color='E5E7EB', width=4)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    for r in p.runs:
                        r.font.size = Pt(9)          # 小五，缩小数据字号减少列内折行
                        r.font.name = '宋体'
                        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 表格固定撑满版心（≈16cm），各列等宽分布
        twips = _compute_autofit_widths(table, _text_width_twips(doc))
        _apply_grid_widths(table, twips)
        for row in table.rows:
            for j, cell in enumerate(row.cells):
                try:
                    cell.width = Cm(twips[j] / 567.0)
                except Exception:
                    pass


# ============================================================
# MD -> Word 引擎（tech / paper 双样式）
# ============================================================
MATH_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'


def _qn_m(tag):
    """math 命名空间下的 clark-notation 标签名（qn 不识别 m: 前缀，故自建）。"""
    return f'{{{MATH_NS}}}{tag}'


def _set_run_font(run, ascii_font, ea_font, size_pt=None, bold=None, color=None):
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), ascii_font)
    rfonts.set(qn('w:hAnsi'), ascii_font)
    rfonts.set(qn('w:eastAsia'), ea_font)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _clear_cell_borders(cell):
    """清除单元格全部边框（实现三线表前必须先去边框）。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tcPr.append(borders)
    else:
        for child in list(borders):
            borders.remove(child)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'none')
        e.set(qn('w:sz'), '0')
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), 'auto')
        borders.append(e)


def _add_row_edge(table, row_idx, edge, color='000000', width=12):
    """给指定行的某条边加单线（三线表用）。"""
    for cell in table.rows[row_idx].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        borders = tcPr.find(qn('w:tcBorders'))
        if borders is None:
            borders = OxmlElement('w:tcBorders')
            tcPr.append(borders)
        e = borders.find(qn(f'w:{edge}'))
        if e is None:
            e = OxmlElement(f'w:{edge}')
            borders.append(e)
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), str(width))
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), color)


def _clear_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            _clear_cell_borders(cell)


def setup_paper_page(doc):
    """A4 纸张 + 页边距 + 页脚居中页码（附件要求：全部用 A4、论文要有页码）。"""
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld1 = OxmlElement('w:fldChar')
    fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    fld2 = OxmlElement('w:fldChar')
    fld2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)


def setup_tech_page(doc):
    """技术文档页面设置：US Letter (12240×15840 twips) +
    上/下 1in、左/右 1.25in、页眉/页脚 0.5in。
    注意：pandoc 默认不写 pgSz/pgMar，故在样式处理最后调用确保写回。"""
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1.25)
    sec.right_margin = Inches(1.25)
    sec.header_distance = Inches(0.5)
    sec.footer_distance = Inches(0.5)


def _set_style_font(style, ascii_font, ea_font, size_pt=None, bold=None):
    """为段落样式（Style 对象）设置中英文字体、字号、加粗。"""
    style.font.name = ascii_font
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), ascii_font)
    rfonts.set(qn('w:hAnsi'), ascii_font)
    rfonts.set(qn('w:eastAsia'), ea_font)
    if size_pt is not None:
        style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold


def setup_paper_fonts(doc):
    """正文=宋体小四(12pt)/1.5倍行距；标题=黑体（03-论文排版与参考文献规范.docx）。
    标题统一用纯黑体（含英文），不加 TNR 混排。"""
    normal = doc.styles['Normal']
    _set_style_font(normal, 'Times New Roman', '宋体')
    normal.font.size = Pt(12)            # 小四
    normal.paragraph_format.line_spacing = 1.5
    for hname, size in (('Heading 1', 15), ('Heading 2', 12),
                        ('Heading 3', 12), ('Heading 4', 12)):
        try:
            st = doc.styles[hname]
            _set_style_font(st, '黑体', '黑体', size_pt=size, bold=True)
            st.paragraph_format.line_spacing = 1.5
        except KeyError:
            pass



def style_tables_threeline(doc):
    """三线表：仅顶线(粗)/栏目线/底线，无竖线、无内部横线；表头黑体不加填充。
    跳过公式编号表（由 number_equations 单独处理）。"""
    for table in doc.tables:
        if _is_equation_table(table):
            continue
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in table.rows:
            for cell in row.cells:
                _clear_cell_borders(cell)
        if len(table.rows) >= 1:
            _add_row_edge(table, 0, 'top', width=12)       # 顶线
            _add_row_edge(table, 0, 'bottom', width=6)     # 栏目线
        if len(table.rows) >= 2:
            _add_row_edge(table, len(table.rows) - 1, 'bottom', width=12)  # 底线
        for cell in table.rows[0].cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    _set_run_font(r, '黑体', '黑体', bold=True, color=RGBColor(0, 0, 0), size_pt=9)
        # 表格表头与数据统一 9pt（小五），减少列内折行
        for i in range(1, len(table.rows)):
            for cell in table.rows[i].cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9)
                        r.font.name = '宋体'
                        r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 表格固定撑满版心（≈16cm），各列等宽分布（与 style_tables 同一算法）
        twips = _compute_autofit_widths(table, _text_width_twips(doc))
        _apply_grid_widths(table, twips)


def style_paper_blocks(doc):
    """首页块（题目/摘要/关键词）+ 参考文献条目字体映射。

    排版规范（来自 03-论文排版与参考文献规范.docx）：
      论文题目/大标题  黑体  小三(15pt)  加粗
      一级标题（一、二、三…）  黑体  小四(12pt)  加粗
      正文            宋体  小四(12pt)  1.5倍行距
      摘要正文        宋体  五号(10.5pt)
      参考文献正文    英文TNR/中文宋体  五号(10.5pt)
      摘要正文        宋体  五号(10.5pt)
    """
    seen_title = False
    in_abstract = False          # 状态：是否处于"摘要"标题后的正文区域
    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue
        style_name = p.style.name if p.style else ''
        # ---- 论文题目：黑体小三加粗居中 ----
        if (not seen_title) and style_name.startswith('Heading 1'):
            for r in p.runs:
                _set_run_font(r, '黑体', '黑体', size_pt=15, bold=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            seen_title = True
            in_abstract = False
        # ---- "摘要"/"Abstract" 标签行：黑体小四 ----
        elif '摘要' in txt or 'Abstract' in txt:
            for r in p.runs:
                _set_run_font(r, '黑体', '黑体', size_pt=12)
            in_abstract = True       # 后续普通段落进入"摘要正文"区
        # ---- "关键词"/"Keywords" 行：宋体五号 ----
        elif re.match(r'^关键词|^Keywords', txt):
            for r in p.runs:
                _set_run_font(r, '宋体', '宋体', size_pt=10.5)   # 五号
            in_abstract = False       # 摘要区域结束
        # ---- 参考文献 [n] 条目：TNR/宋体五号 ----
        elif re.match(r'^\[\d+\]', txt):
            for r in p.runs:
                _set_run_font(r, 'Times New Roman', '宋体', size_pt=10.5)
            in_abstract = False
        # ---- 遇到其他标题 → 退出摘要区 ----
        elif style_name.startswith('Heading'):
            in_abstract = False
        # ---- 摘要正文段落：宋体五号（区别于正文的宋体小四）----
        elif in_abstract:
            for r in p.runs:
                _set_run_font(r, '宋体', '宋体', size_pt=10.5)


def _text_width_twips(doc):
    """返回当前节的版心宽度（twips）。两种样式最终均为 A4 + 2.5cm 边距，
    故直接计算：21cm - 2*2.5cm = 16cm ≈ 9071 twips。"""
    try:
        from docx.shared import Cm
        emu = (Cm(21) - 2 * Cm(2.5))
        return int(emu / 635)   # 1 twip = 635 EMU
    except Exception:
        return 9071


def _set_math_size(eq_om, half_pt):
    """把公式内所有数学 run（m:r）的字号统一设为 half_pt（半磅，如 21=10.5pt）。
    某些极宽公式在本项目的论文版（正文 12pt）下会超出表格左列而折行，故统一
    降到五号 10.5pt 显示，既符合论文显示公式习惯，也保证最宽公式（如 (6)）单行。"""
    for mr in eq_om.iter(_qn_m('r')):
        rpr = mr.find(_qn_m('rPr'))
        if rpr is None:
            rpr = OxmlElement('m:rPr'); mr.insert(0, rpr)
        sz = rpr.find(qn('w:sz'))
        if sz is None:
            sz = OxmlElement('w:sz'); rpr.append(sz)
        sz.set(qn('w:val'), str(half_pt))


def _build_eq_table(eq_om, num_text, eq_w, num_w, eq_sz=21):
    """构建无框 1×2 表格：左列(eq_w)放公式居中，右列(num_w)放编号右对齐。
    eq_om 为已拆包的行内 <m:oMath> 节点，本函数将其**移动**入左列（不深拷贝），
    规避历史「公式变空白框」的渲染失败。编号恒在页面最右固定列，绝不随公式宽度
    重叠——彻底解决「编号离宽公式太近」问题。"""
    def _nil_borders(parent, edges=('top', 'left', 'bottom', 'right',
                                    'insideH', 'insideV')):
        b = OxmlElement('w:tcBorders')
        for edge in edges:
            e = OxmlElement(f'w:{edge}')
            e.set(qn('w:val'), 'none'); e.set(qn('w:sz'), '0')
            e.set(qn('w:space'), '0'); e.set(qn('w:color'), 'auto')
            b.append(e)
        parent.append(b)

    def _zero_margin(parent, edges=('top', 'left', 'bottom', 'right')):
        m = OxmlElement('w:tcMar')
        for edge in edges:
            e = OxmlElement(f'w:{edge}'); e.set(qn('w:w'), '0')
            e.set(qn('w:type'), 'dxa'); m.append(e)
        parent.append(m)

    tbl = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    _nil_borders(tblPr)
    tblW = OxmlElement('w:tblW'); tblW.set(qn('w:w'), str(eq_w + num_w))
    tblW.set(qn('w:type'), 'dxa'); tblPr.append(tblW)
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'autofit')
    tblPr.append(layout)
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'left'); tblPr.append(jc)
    _zero_margin(tblPr)
    tbl.append(tblPr)
    grid = OxmlElement('w:tblGrid')
    g1 = OxmlElement('w:gridCol'); g1.set(qn('w:w'), str(eq_w)); grid.append(g1)
    g2 = OxmlElement('w:gridCol'); g2.set(qn('w:w'), str(num_w)); grid.append(g2)
    tbl.append(grid)
    tr = OxmlElement('w:tr'); tbl.append(tr)
    # 左单元格：公式居中
    tc1 = OxmlElement('w:tc')
    tcPr1 = OxmlElement('w:tcPr')
    tcW1 = OxmlElement('w:tcW'); tcW1.set(qn('w:w'), str(eq_w))
    tcW1.set(qn('w:type'), 'dxa'); tcPr1.append(tcW1)
    _nil_borders(tcPr1); _zero_margin(tcPr1)
    v1 = OxmlElement('w:vAlign'); v1.set(qn('w:val'), 'center'); tcPr1.append(v1)
    tc1.append(tcPr1)
    p1 = OxmlElement('w:p'); pPr1 = OxmlElement('w:pPr')
    jc1 = OxmlElement('w:jc'); jc1.set(qn('w:val'), 'center'); pPr1.append(jc1)
    _set_math_size(eq_om, eq_sz)
    p1.append(pPr1); p1.append(eq_om)
    tc1.append(p1); tr.append(tc1)
    # 右单元格：编号右对齐
    tc2 = OxmlElement('w:tc')
    tcPr2 = OxmlElement('w:tcPr')
    tcW2 = OxmlElement('w:tcW'); tcW2.set(qn('w:w'), str(num_w))
    tcW2.set(qn('w:type'), 'dxa'); tcPr2.append(tcW2)
    _nil_borders(tcPr2); _zero_margin(tcPr2)
    v2 = OxmlElement('w:vAlign'); v2.set(qn('w:val'), 'center'); tcPr2.append(v2)
    tc2.append(tcPr2)
    p2 = OxmlElement('w:p'); pPr2 = OxmlElement('w:pPr')
    jc2 = OxmlElement('w:jc'); jc2.set(qn('w:val'), 'right'); pPr2.append(jc2)
    p2.append(pPr2)
    r2 = OxmlElement('w:r'); rPr2 = OxmlElement('w:rPr')
    sz2 = OxmlElement('w:sz'); sz2.set(qn('w:val'), str(eq_sz)); rPr2.append(sz2)
    r2.append(rPr2)
    t2 = OxmlElement('w:t')
    t2.set(qn('xml:space'), 'preserve'); t2.text = num_text
    r2.append(t2); p2.append(r2)
    tc2.append(p2); tr.append(tc2)
    return tbl


def number_equations(doc, eq_tags=None, log_cb=None):
    """为 md 中标记了 \tag{N} 的独立公式，在 Word 里于同行右侧显示 (N)。

    编号来源：md 源文件（经 extract_eq_tags 提取，与 $$ 块顺序一致），
    保证「md 里的编号 == Word 里的编号」。pandoc 不支持 \tag，故转换前已剔除，
    这里仅按 md 给的号落到 Word。

    实现：把每个独立公式从块级 <m:oMathPara> 就地拆包为行内 <m:oMath>（移动节点、
    不深拷贝，规避历史「公式变空白框」渲染失败），再整体包进一个**无框 1×2 表格**：
      · 左列（版心宽 - 编号列宽）：公式居中——即使极宽公式也只在此列内居中，
        绝不会与右列编号重叠；
      · 右列（固定约 0.6in）：编号 (N) 右对齐、顶到同行最右。
    这样既解决「编号离公式太近」，又保证编号始终在页面最右、与公式同行。"""
    def log(m):
        if log_cb:
            log_cb(m)
    try:
        text_w = _text_width_twips(doc)
        NUM_W = 380                      # 编号列宽 ~0.27in，仅容纳 "(99)"
        EQ_W = text_w - NUM_W            # 公式列宽（≈版心全宽）
        EQ_SZ = 20                       # 公式显示字号 10pt，保证最宽公式在 Word 中单行
        tags = list(eq_tags) if eq_tags else []
        auto = (len(tags) == 0)          # 无 tag 信息时退化为从 1 自动编号（兼容旧调用）
        n = 0
        for p in list(doc.paragraphs):
            # 1) 取出公式节点（优先 oMathPara 内包裹的 oMath，否则行内 oMath）
            op = p._p.find('.//%s' % _qn_m('oMathPara'))
            if op is not None:
                om = op.find(_qn_m('oMath'))
                if om is None:
                    continue
                # 就地「拆包」：把块级 oMathPara 换成行内 oMath（移动节点，不深拷贝）
                op.addprevious(om)
                op.getparent().remove(op)
                eq = om
            else:
                eq = p._p.find('.//%s' % _qn_m('oMath'))
                if eq is None:
                    continue
            # 2) 跳过含其它文字的段落（行内公式不编号）
            if ''.join(r.text for r in p.runs).strip():
                continue
            # 3) 取编号：与 $$ 块顺序一致的 eq_tags；无 tag 则不编号
            if auto:
                n += 1
                num = n
            else:
                if not tags:
                    break
                num = tags.pop(0)
                if num is None:
                    continue   # md 中未标记编号的公式，跳过
            # 4) 构建无框 1×2 表格：[公式居中 | (N) 右对齐]
            tbl = _build_eq_table(eq, f'({num})', EQ_W, NUM_W, EQ_SZ)
            # 5) 用表格替换原公式段落（原段落在 unpack 后已无 oMath，仅作占位被替换）
            p._p.addprevious(tbl)
            p._p.getparent().remove(p._p)
        log("  公式编号处理完成（无框表格：公式居中 + 编号同行右侧，取自 md 的 \tag）")
    except Exception as e:
        log(f"  公式编号处理跳过: {e}")




def _style_image_captions(doc):
    """pandoc 将图片 alt 文字输出为 style='ImageCaption' 的段落（紧跟在图片段之后）。
    本函数将其设为：居中对齐 + 宋体五号(10.5pt) —— 与摘要/参考文献同档（规范未规定图名
    字号，取论文常规的「辅助小字=五号」口径，最稳妥满足格式要求；如需小五(9pt)改 size_pt）。"""
    for p in doc.paragraphs:
        if p.style and p.style.name in ('ImageCaption', 'Image Caption'):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                _set_run_font(r, '宋体', '宋体', size_pt=10.5)


def _apply_shared_style(doc, table_styler, log_cb=None, eq_tags=None):
    """MD→Word（tech）与 论文版（paper）的公共排版样式。

    设计目标（用户 7/15 要求）：让两种模式的**唯一差异**收敛到表格表头。
    因此除 table_styler 外，页面(A4+页码)、字体(宋体正文小四/黑体标题)、
    首页块(题目居中加粗/摘要五号/参考文献)、公式处理全部一致。

    table_styler：唯一差异点，传入不同的表格样式函数：
      · style_tables            -> 深蓝表头 + 灰边（MD→Word）
      · style_tables_threeline  -> 三线表（论文版）
    """
    def log(m):
        if log_cb:
            log_cb(m)

    setup_paper_fonts(doc)          # 字体：宋体正文小四/1.5倍、黑体标题加粗
    table_styler(doc)               # ← 唯一差异：表格表头样式
    style_paper_blocks(doc)         # 首页块：题目居中加粗/摘要五号/参考文献五号
    number_equations(doc, eq_tags, log_cb)   # 公式编号取自 md 的 \tag，落到同行右侧
    _style_image_captions(doc)                  # 图片标题居中 + 小五号宋体
    # 页面设置（A4/页码）必须最后做：前面的表格/公式处理可能重建 body 结构，
    # 导致 sectPr 的 pgSz/pgMar 被清除，故放在末尾确保最终写回。
    setup_paper_page(doc)


def apply_paper_style(doc, log_cb=None, eq_tags=None):
    def log(m):
        if log_cb:
            log_cb(m)

    _apply_shared_style(doc, style_tables_threeline, log_cb, eq_tags)
    log("  论文格式样式已应用（A4/页码/宋体/三线表/黑体标题/摘要五号）")
    log("  ⚠️ 论文合规检查清单：请确保无封面、无个人照片、任何人名（含指导老师/"
        "专家/采访人）、无“鸣谢”；参考文献≥6篇（英文≥2）、网页/报刊不可列入；"
        "表名在上、图名在下、单位用国际单位。")


def extract_eq_tags(text):
    """按文档顺序扫描所有 $$...$$ 显示公式块，返回编号列表。

    - 块内若含 \\tag{N} / \\tag(N)（单或双反斜杠均可）则记录其数字 N；
    - 否则记录 None（该块为无编号公式）。
    返回的列表与 $$ 块一一对应，供 number_equations 按相同顺序把编号落到 Word。"""
    tags = []
    for m in re.finditer(r'\$\$(.*?)\$\$', text, flags=re.DOTALL):
        body = m.group(1)
        tm = re.search(r'\\{1,2}tag\s*[({]\s*(\d+)\s*[)}]', body)
        tags.append(int(tm.group(1)) if tm else None)
    return tags


def _strip_unsupported_math(text):
    """剔除 pandoc texmath 不支持的 \\tag / \\label / \\nonumber 等命令，
    否则整块 $$...$$ 公式会被 pandoc 降级为纯文本（表现为「公式与编号都没了」）。

    注意：源文件常见笔误把 \\tag 写成 \\\\tag（双反斜杠），故按 1~2 个前导反斜杠匹配；
    同时兼容 \\tag{N} 与 \\tag(N) 两种写法。只作用于 $$...$$ 块内部，不触碰正文其它反斜杠。"""
    def clean(body):
        body = re.sub(r'\\{1,2}tag\s*[({][^)}]*[)}]', '', body)
        body = re.sub(r'\\{1,2}label\s*\{[^}]*\}', '', body)
        body = re.sub(r'\\{1,2}(?:nonumber|allowbreak)\b', '', body)
        return body
    return re.sub(r'\$\$(.*?)\$\$',
                  lambda m: '$$' + clean(m.group(1)) + '$$',
                  text, flags=re.DOTALL)


def preprocess_md(text):
    """将 HTML <img> 标签转换为 pandoc 可识别的 Markdown 图片语法，
    使 docx 导出也能嵌入图片；同时剔除浮动清理用的 clear:both div（docx 不需要），
    并去掉 pandoc 无法转换的 \\tag 等数学命令。源 .md 文件不会被修改。"""
    def repl_img(m):
        attrs = m.group(1)
        src = re.search(r'src\s*=\s*"([^"]*)"', attrs)
        if not src:
            return ''
        alt = re.search(r'alt\s*=\s*"([^"]*)"', attrs)
        alt_text = alt.group(1) if alt else ''
        img = f'![{alt_text}]({src.group(1)})'
        width = re.search(r'width\s*=\s*"(\d+)"', attrs)
        if width:
            # 假设 96 DPI：px -> inch，pandoc 用 inch 指定图片宽度
            inch = int(width.group(1)) / 96.0
            img += f'{{width={inch:.2f}in}}'
        return img
    text = re.sub(r'<img\b([^>]*)>', repl_img, text, flags=re.IGNORECASE)
    text = re.sub(r'<div[^>]*clear\s*:\s*both[^>]*>\s*</div>', '', text,
                  flags=re.IGNORECASE)

    # ---- <sup>/<sub> HTML 标签 → pandoc 原生上标/下标语法 ----
    # 背景：pandoc 的 markdown reader 把 <sup>..</sup> 当作 raw_html 内联，
    # docx writer 会直接丢弃 raw html（仅保留其中文字），导致上标失效。
    # pandoc markdown 默认启用 superscript(^x^)/subscript(~x~) 扩展，转成原生语法后
    # docx 才会渲染为真正的上标/下标。上标内容不能含裸空格，需转义为 "\ "。
    # 典型用途：正文引用编号如 <sup>[3,5]</sup> → ^[3,5]^（Word 中显示为上标）。
    def _repl_sup(m):
        return '^' + m.group(1).replace(' ', r'\ ') + '^'
    def _repl_sub(m):
        return '~' + m.group(1).replace(' ', r'\ ') + '~'
    text = re.sub(r'<sup>(.+?)</sup>', _repl_sup, text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r'<sub>(.+?)</sub>', _repl_sub, text, flags=re.IGNORECASE | re.DOTALL)

    text = _strip_unsupported_math(text)

    # ---- 修复 pandoc 不识别的"段落后紧跟列表"和"连续块引用" ----
    # pandoc 要求列表前必须有空行，否则把 "- item" 当作段落内文字吞掉；
    # 同理，连续两行 > 块引用（无空行）会被合并为一段。
    # 本步骤在转换管道内自动插入空行，不改 md 源文件。
    text = _fix_pandoc_paragraph_breaks(text)

    return text


def _fix_pandoc_paragraph_breaks(text):
    """在送入 pandoc 前修复两种常见换行丢失问题：

    1）普通段落后紧跟列表行（-/*/+/数字.）→ 中间插空行；
    2）连续 > 块引用行（无空行）→ 之间插空行使各成独立段落。
    跳过代码围栏（```）、表格行（|）、标题（#）等特殊块。"""
    LIST_RE = re.compile(r'^(\s*)([-*+]|\d+\.)\s+')
    lines = text.split('\n')
    out = []
    in_code = False
    for i, line in enumerate(lines):
        stripped = line.rstrip()
        # Track code fence
        if stripped.startswith('```'):
            in_code = not in_code
            out.append(line)
            continue
        if in_code:
            out.append(line)
            continue

        # Check if THIS line starts a list / blockquote
        is_list = bool(LIST_RE.match(stripped))
        is_bq = stripped.startswith('> ')

        if is_list or is_bq:
            # Look back: find the most recent non-empty line that was added
            need_blank = False
            for j in range(len(out) - 1, -1, -1):
                prev = out[j].rstrip()
                if prev == '':
                    break                    # already has blank line → OK
                if prev.startswith('```'):
                    break                    # code fence edge → don't touch
                if LIST_RE.match(prev.lstrip()):
                    break                    # previous is list item → continuation OK
                if is_bq and prev.startswith('> '):
                    need_blank = True       # consecutive BQ → insert blank
                    break
                if is_list:
                    need_blank = True        # paragraph before list → insert blank
                    break
                if prev.startswith('|'):
                    need_blank = True        # table row before list/bq
                    break
                if prev.startswith('#'):
                    break                    # heading before list → OK
                break                        # other case (empty, etc.)
            if need_blank:
                out.append('')              # insert blank line

        out.append(line)

    return '\n'.join(out)


def renumber_equations_md(md_path, all_blocks=True, log_cb=None):
    """重排 md 源文件中公式的 \\tag 编号（用户需求 2：代码修改导致公式增减后，
    对应编号需重新排列）。

    - 扫描所有 $$...$$ 显示公式块，按文档顺序重新编号为 \\tag{1}、\\tag{2}、…
    - 默认 all_blocks=True：给【每一个】 $$ 块都补上顺序编号（即所有显示公式
      都编号）；若某块已有 \\tag，则替换为新序号。
    - all_blocks=False（--tagged-only）：仅重排「原本就带 \\tag」的块，未标记的
      保持无编号。
    - 同时把笔误 \\tag（双反斜杠）规范为单反斜杠 \\tag；编号统一写在块末。
    源 .md 文件会被【原地修改】。"""
    def log(m):
        if log_cb:
            log_cb(m)
    try:
        with open(md_path, encoding='utf-8') as _f:
            text = _f.read()
    except Exception as e:
        log(f"错误: 无法读取 {md_path}: {e}")
        return False

    counter = 0

    def repl(m):
        nonlocal counter
        body = m.group(1)
        had_tag = re.search(r'\\{1,2}tag', body) is not None
        # 先移除块内已有的 tag/label，再按需重新追加
        cleaned = re.sub(r'\\{1,2}tag\s*[({][^)}]*[)}]', '', body)
        cleaned = re.sub(r'\\{1,2}label\s*\{[^}]*\}', '', cleaned)
        cleaned = re.sub(r'\\{1,2}(?:nonumber|allowbreak)\b', '', cleaned)
        if all_blocks or had_tag:
            counter += 1
            # 规范化：编号写在块末（去掉尾部空白后追加），单反斜杠 + 花括号
            cleaned = cleaned.rstrip() + f' \\tag{{{counter}}}'
            return '$$' + cleaned + '$$'
        return '$$' + cleaned + '$$'

    new_text = re.sub(r'\$\$(.*?)\$\$', repl, text, flags=re.DOTALL)
    if new_text == text:
        log(f"  无需重排（{os.path.basename(md_path)} 的公式编号已连续）")
        return True
    with open(md_path, 'w', encoding='utf-8') as _f:
        _f.write(new_text)
    log(f"  已重排 {counter} 个公式编号，写入 {os.path.basename(md_path)}"
        f"（现在用 skill 转换即可与 Word 保持一致）")
    return True


def convert_file(md_path, docx_path, resource_path=None, style='tech', log_cb=None):
    def log(msg):
        if log_cb:
            log_cb(msg)
        else:
            print(msg)

    pandoc = find_pandoc()
    if not pandoc:
        log("错误: 未找到 pandoc，请先安装并加入 PATH")
        return False

    if not os.path.isfile(md_path):
        log(f"错误: 文件不存在: {md_path}")
        return False

    md_dir = os.path.dirname(os.path.abspath(md_path))
    res_path = resource_path or md_dir

    # 预处理：把 HTML <img> 标签转为 Markdown 图片语法，使 docx 也能嵌入图片；
    # 同时剔除浮动清理用的 clear:both div（docx 不需要）。源 .md 文件不改。
    with open(md_path, encoding='utf-8') as _f:
        _raw = _f.read()

    # Word 对原生 SVG 支持很差（常常显示空白），提前给出明确警告
    if re.search(r'\.svg(?:\s|\)|"|\s*\])', _raw, re.IGNORECASE):
        log("⚠️ 警告: 检测到 .svg 图片引用，Word 可能无法渲染（显示空白）。"
            "建议先转为 .png 再引用（参考 肌电算法说明.md 的做法）。")
    # 论文模式：通栏图上限 130mm ≈ 5.12in，超出给警告
    if style == 'paper':
        for m in re.finditer(r'\{width=([\d.]+)in\}', _raw):
            inch = float(m.group(1))
            if inch > 5.12:
                log(f"  ⚠️ 论文模式：图片宽度 {inch:.2f}in 超过通栏图上限 "
                    f"(≈130mm)，建议缩小。")
    _proc = preprocess_md(_raw)
    # 提取 md 中 $$ 块的 \tag 编号（与块顺序一致），交给 number_equations 原样落到 Word
    eq_tags = extract_eq_tags(_raw)
    _tmp = tempfile.NamedTemporaryFile(mode='w', suffix='.md', encoding='utf-8',
                                       delete=False, dir=md_dir)
    _tmp.write(_proc)
    _tmp.close()
    _input_md = _tmp.name
    try:
        pandoc_args = [
            # -inline_notes: 禁用 ^[...] 内联脚注语法，否则正文上标引用
            # ^[3,5]^（由 <sup>[3,5]</sup> 转来）会被误解析为脚注而非上标。
            '-f', 'markdown-citations-inline_notes+tex_math_dollars',
            '-t', 'docx',
            '--wrap=preserve',
            '--mathml',
            '--resource-path', res_path,
            '-o', docx_path,
            _input_md,
        ]

        log(f"正在转换: {os.path.basename(md_path)}")
        proc = subprocess.run([pandoc] + pandoc_args, capture_output=True)
        _stderr = (proc.stderr.decode('utf-8', errors='replace')
                   if isinstance(proc.stderr, (bytes, bytearray)) else (proc.stderr or ''))
        if proc.returncode != 0:
            log(f"pandoc 转换失败:\n{_stderr}")
            return False

        if not os.path.isfile(docx_path):
            log("错误: 转换后未找到输出文件")
            return False

        try:
            doc = Document(docx_path)
            if style == 'paper':
                apply_paper_style(doc, log, eq_tags)
            else:
                # MD→Word（tech）：与论文版共用全部排版，唯一差异=深蓝表头表格
                _apply_shared_style(doc, style_tables, log, eq_tags)
                log("  样式已应用（与论文版一致，仅表格为深蓝表头）")
            doc.save(docx_path)
        except Exception as e:
            log(f"  后处理样式应用失败: {e}")
            return False

        log(f"  完成 -> {os.path.basename(docx_path)}")
        return True
    finally:
        try:
            os.unlink(_input_md)
        except OSError:
            pass



# ============================================================
# CLI 子命令实现
# ============================================================
def _cli_out_name(md_path, style):
    """批量模式下按 style 从输入文件名派生输出 .docx 名。"""
    base = os.path.splitext(os.path.basename(md_path))[0]
    return base + ("_论文版.docx" if style == "paper" else ".docx")


def _expand_md_inputs(inputs, log):
    """把 CLI input 列表（文件 / 目录 / glob 通配）展开为去重保序的 .md 文件列表。
    目录只扫单层 *.md（与 GUI 批量一致，不递归子目录）。"""
    md_files = []
    for item in inputs:
        if os.path.isdir(item):
            md_files.extend(sorted(glob.glob(os.path.join(item, "*.md"))))
        else:
            matched = sorted(glob.glob(item))
            if matched:
                md_files.extend(matched)
            elif os.path.exists(item):
                md_files.append(item)
            else:
                log(f"⚠ 未找到输入: {item}")
    seen = set()
    return [f for f in md_files if not (f in seen or seen.add(f))]


def run_to_docx(args, log):
    inputs = args.input if isinstance(args.input, list) else [args.input]
    md_files = _expand_md_inputs(inputs, log)
    if not md_files:
        log("未找到任何 .md 文件")
        return False

    out = args.output
    # 单文件且 -o 不是目录 -> 保持旧行为：-o 即输出文件路径
    out_is_dir = os.path.isdir(out) or out.endswith(("/", "\\"))
    if len(md_files) == 1 and not out_is_dir:
        return convert_file(md_files[0], out, args.resource_path, args.style, log)

    # 批量：-o 视为输出目录，逐个派生同名 .docx
    # 防呆：批量下若 -o 误写成 .docx 文件名（且非已有目录），改用其父目录，避免建出 “xxx.docx” 文件夹
    if not os.path.isdir(out) and out.lower().endswith(".docx"):
        parent = os.path.dirname(out) or "."
        log(f"⚠ 批量模式下 -o 应为【目录】，但收到文件名 “{out}”；已改用其所在目录 “{parent}”。")
        out = parent
    os.makedirs(out, exist_ok=True)
    log(f"批量模式：共 {len(md_files)} 个文件 -> {os.path.abspath(out)}")
    all_ok = True
    for f in md_files:
        outp = os.path.join(out, _cli_out_name(f, args.style))
        log(f"\n--- {os.path.basename(f)} -> {os.path.basename(outp)} ---")
        try:
            if convert_file(f, outp, args.resource_path, args.style, log) is False:
                all_ok = False
        except Exception as e:
            log(f"  ⚠ 失败: {e}")
            all_ok = False
    return all_ok




def build_cli():
    p = argparse.ArgumentParser(prog="docx-toolkit.pyw",
                                 description="Markdown -> Word 文档转换工具集")
    sub = p.add_subparsers(dest="cmd", required=True)

    sp = sub.add_parser("to-docx", help="Markdown -> Word（支持单文件 / 多文件 / 目录批量）")
    sp.add_argument("input", nargs="+",
                    help="一个或多个 .md 文件、目录（扫单层 *.md）或 glob 通配")
    sp.add_argument("-o", "--output", required=True,
                    help="单文件时为输出 .docx 路径；多文件/目录时为输出目录")
    sp.add_argument("--style", choices=["tech", "paper"], default="tech")
    sp.add_argument("--resource-path")
    sp.set_defaults(func=run_to_docx)

    rp = sub.add_parser("renumber-equations",
                        help="重排 md 源文件里公式的 \\tag 编号（公式增减后调用）")
    rp.add_argument("input", help="要重排编号的 .md 文件")
    rp.add_argument("--tagged-only", action="store_true",
                    help="只重排原本带 \\tag 的块；默认给所有 $$ 块顺序编号")
    rp.set_defaults(func=run_renumber)


    return p


def run_renumber(args, log):
    return renumber_equations_md(args.input,
                                  all_blocks=not args.tagged_only,
                                  log_cb=log)


def cli_main():
    p = build_cli()
    args = p.parse_args()
    ok = args.func(args, print)
    sys.exit(0 if (ok is not False) else 1)


# ============================================================
# GUI
# ============================================================

# 功能分组定义：(组标题, [(命令值, 标签), ...])
# 第一行两种 MD→Word 输出格式
FUNC_GROUPS = [
    ("MD 文档", [
        ("md-word",   "MD → Word（深蓝网格表）"),
        ("md-paper",  "MD → Word 论文版（三线表）"),
    ]),
]

LOCK_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".toolkit.lock")


def _ensure_single_instance():
    """确保只有一个 docx-toolkit.pyw 在运行。读取 lock 文件中的 PID，若存在且进程仍存活则终止它。"""
    try:
        if os.path.exists(LOCK_FILE):
            old_pid_str = open(LOCK_FILE, encoding="utf-8").read().strip()
            if old_pid_str:
                import signal
                old_pid = int(old_pid_str)
                # 尝试终止旧进程（仅限同用户、同名进程）
                try:
                    os.kill(old_pid, signal.SIGTERM)
                    import time; time.sleep(0.3)
                except ProcessLookupError:
                    pass  # 进程已不存在，直接覆盖
        # 写入当前 PID
        with open(LOCK_FILE, "w", encoding="utf-8") as f:
            f.write(str(os.getpid()))
    except Exception:
        pass


def _cleanup_lock():
    """退出时清理锁文件。"""
    try:
        if os.path.exists(LOCK_FILE):
            os.remove(LOCK_FILE)
    except Exception:
        pass


def run_gui():
    import tkinter as tk
    from tkinter import ttk, filedialog, messagebox

    # ---- 单实例保证 ----
    _ensure_single_instance()

    root = tk.Tk()
    root.title("docx-toolkit — Markdown 文档转换工具")
    root.geometry("940x620")
    root.protocol("WM_DELETE_WINDOW", lambda: (_cleanup_lock(), root.destroy()))

    mode_var = tk.StringVar(value="md-word")
    batch_var = tk.BooleanVar(value=False)          # 单/批量复选

    main = ttk.Frame(root, padding=10)
    main.pack(fill=tk.BOTH, expand=True)

    # ======== 辅助函数（先定义，供下面控件 command 引用） ========
    def _derive_output(inp, mode):
        """按功能值与输入路径推导默认输出路径。
        批量模式(mode 以 batch- 开头)时返回目录本身；单文件模式返回 .docx 文件名。"""
        if mode.startswith("batch-"):
            # 输入是目录就直接用该目录作为默认输出位置
            return inp if os.path.isdir(inp) else os.path.dirname(inp) or "."
        base = os.path.splitext(inp)[0]
        if mode == "md-word":    return base + ".docx"
        if mode == "md-paper":   return base + "_论文版.docx"
        return base + ".docx"

    def _effective_mode(mode, is_batch):
        """批量复选时把单文件功能映射为批量标记（batch- 前缀）。"""
        if not is_batch:
            return mode
        mapping = {
            "md-word":   "batch-md-word",
            "md-paper":  "batch-md-paper",
        }
        return mapping.get(mode, mode)

    def _on_mode_change():
        """功能或批量复选变化时更新输入框的浏览类型提示。"""
        mode = mode_var.get()
        is_batch = batch_var.get()
        eff = _effective_mode(mode, is_batch)
        inp = in_e.get().strip()
        if inp:
            out_e.delete(0, tk.END)
            out_e.insert(0, _derive_output(inp, eff))

    # ======== 第一行：功能单选（MD→Word / MD→论文版 两种） ========
    ttk.Label(main, text="输出格式：", font=("", 10, "bold")).pack(anchor=tk.W, pady=(4, 0))
    row1 = ttk.Frame(main)
    row1.pack(fill=tk.X, pady=4)
    for gtitle, items in FUNC_GROUPS:
        gf = ttk.LabelFrame(row1, text=gtitle)
        gf.pack(side=tk.LEFT, padx=6, fill=tk.X, expand=True)
        for v, label in items:
            ttk.Radiobutton(gf, text=label, variable=mode_var, value=v,
                            command=_on_mode_change).pack(side=tk.LEFT, padx=4, pady=2)

    ttk.Label(main, text="两种格式排版完全一致，唯一区别是表格样式（tech=深蓝网格表 / paper=三线表）。",
              font=("", 9), foreground="#555555").pack(anchor=tk.W, pady=(0, 2))

    # ======== 第二行：批量模式复选框（与第一行复选关系） ========
    row2 = ttk.Frame(main)
    row2.pack(fill=tk.X, pady=(2, 0))
    ttk.Checkbutton(row2, text="批量模式（勾选后对目录下所有 .md 文件转换）",
                    variable=batch_var, command=_on_mode_change).pack(side=tk.LEFT)

    # ======== 输入路径 ========
    ttk.Label(main, text="输入路径:", font=("", 10, "bold")).pack(anchor=tk.W, pady=(6, 0))
    inf = ttk.Frame(main)
    inf.pack(fill=tk.X, pady=4)
    in_e = ttk.Entry(inf)
    in_e.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 5))

    def browse_in():
        mode = mode_var.get()
        is_batch = batch_var.get()
        # 批量模式 -> 选目录；否则选单个 .md 文件
        if is_batch:
            p = filedialog.askdirectory()
        else:  # 两种 MD 功能（md-word/md-paper）均为 .md 源
            p = filedialog.askopenfilename(filetypes=[("Markdown", "*.md")])
        if p:
            in_e.delete(0, tk.END)
            in_e.insert(0, p)
            out_e.delete(0, tk.END)
            out_e.insert(0, _derive_output(p, _effective_mode(mode, is_batch)))
    ttk.Button(inf, text="浏览...", command=browse_in).pack(side=tk.RIGHT)

    # ======== 输出路径 ========
    ttk.Label(main, text="输出路径:", font=("", 10, "bold")).pack(anchor=tk.W, pady=(6, 0))
    outf = ttk.Frame(main)
    outf.pack(fill=tk.X, pady=4)
    out_e = ttk.Entry(outf)
    out_e.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 5))

    def browse_out():
        mode = mode_var.get()
        is_batch = batch_var.get()
        eff = _effective_mode(mode, is_batch)
        # 批量 / 批量映射 -> 选目录；单文件 MD 转换 -> 保存 .docx
        if is_batch or eff.startswith("batch"):
            p = filedialog.askdirectory()
        else:  # md-word / md-paper
            p = filedialog.asksaveasfilename(defaultextension=".docx")
        if p:
            out_e.delete(0, tk.END)
            out_e.insert(0, p)
    ttk.Button(outf, text="浏览...", command=browse_out).pack(side=tk.RIGHT)

    # ======== 日志 ========
    log_text = tk.Text(main, height=8, wrap=tk.WORD)
    log_text.pack(fill=tk.BOTH, expand=True, pady=(8, 0))

    def log(m):
        log_text.insert(tk.END, m + "\n")
        log_text.see(tk.END)
        root.update()

    # ======== 开始转换逻辑 ========
    def _run_single(mode, inp, outp, res, log):
        """单文件转换分发（md-word / md-paper 两种功能）。"""
        if mode == "md-word":
            convert_file(inp, outp, res, 'tech', log)
        elif mode == "md-paper":
            convert_file(inp, outp, res, 'paper', log)
        else:
            log(f"未知功能: {mode}")

    def _run_batch(mode, target_dir, out_dir, res, log):
        """批量模式：遍历目录下所有 .md 文件逐个调 _run_single。"""
        # 防呆：批量下若输出路径误写成 .docx 文件名（且非已有目录），改用其父目录
        if not os.path.isdir(out_dir) and out_dir.lower().endswith(".docx"):
            parent = os.path.dirname(out_dir) or "."
            log(f"⚠ 批量模式下输出路径应为【目录】，但收到文件名 “{out_dir}”；已改用 “{parent}”。")
            out_dir = parent

        pat = "*.md"
        files = sorted(glob.glob(os.path.join(target_dir, pat)))
        if not files:
            log(f"未找到文件: {target_dir}/{pat}")
            return
        log(f"批量模式：共 {len(files)} 个文件 -> {os.path.abspath(out_dir)}")
        for f in files:
            out_name = os.path.basename(_derive_output(f, mode))
            outp = os.path.join(out_dir, out_name)
            log(f"\n--- {os.path.basename(f)} -> {out_name} ---")
            try:
                _run_single(mode, f, outp, res, log)
            except Exception as e:
                log(f"  ⚠ 失败: {e}")

    def start():
        mode = mode_var.get()
        is_batch = batch_var.get()
        log_text.delete(1.0, tk.END)
        inp = in_e.get().strip()
        if not inp:
            messagebox.showwarning("提示", "请选择输入路径")
            return
        res = None  # 图片已内联进 md (data URI)，无需外部资源路径
        try:
            if is_batch:
                target_dir = inp if os.path.isdir(inp) else os.path.dirname(inp)
                out_dir = out_e.get().strip() or target_dir
                _run_batch(mode, target_dir, out_dir, res, log)
            else:
                outp = out_e.get().strip() or _derive_output(inp, mode)
                _run_single(mode, inp, outp, res, log)
        except Exception as e:
            log(f"\n错误: {e}")
        else:
            log("\n✅ 完成")

    # ======== 底部按钮 ========
    bf = ttk.Frame(main)
    bf.pack(fill=tk.X, pady=8)
    ttk.Button(bf, text="开始转换", command=start, width=20).pack(side=tk.LEFT, padx=4)
    def do_renumber():
        inp = in_e.get().strip()
        if not inp or not os.path.isfile(inp):
            messagebox.showwarning("提示", "请先在「输入路径」选择要重排编号的 .md 文件")
            return
        log_text.delete(1.0, tk.END)
        log("重排公式编号（将修改 md 源文件）…")
        try:
            renumber_equations_md(inp, all_blocks=True, log_cb=log)
        except Exception as e:
            log(f"错误: {e}")
        else:
            log("\n✅ 完成。之后用「开始转换」即可让 Word 编号与 md 一致。")
    ttk.Button(bf, text="重排公式编号", command=do_renumber, width=20).pack(side=tk.LEFT, padx=4)
    root.mainloop()
    _cleanup_lock()


def main():
    if len(sys.argv) > 1:
        cli_main()
    else:
        run_gui()


if __name__ == "__main__":
    main()
