# -*- coding: utf-8 -*-
"""
docx_utils.pyw - docx 与 Markdown 互转 / 报告生成 / 批量转换 / docx 转 PDF

通用工具, 不绑定任何具体项目。基于 pandoc + python-docx, docx 转 PDF 需 LibreOffice。

双模式:
  - 双击 / 无参数  -> GUI 模式(tkinter 文件/文件夹选择)
  - 指令模式       -> 加 --cli 或直接以子命令开头, 例如:
      pythonw docx_utils.pyw --cli to-docx a.md -o a.docx
      pythonw docx_utils.pyw to-docx a.md -o a.docx        (自动识别为 CLI)

子命令:
  to-docx      Markdown -> docx
  to-md        docx -> Markdown
  gen-report   Markdown -> 正式报告(docx, 含目录/样式/页码)
  batch        目录批量互转 (md2docx / docx2md)
  docx2pdf     docx -> PDF (需 LibreOffice)
  package      目录内 Markdown 打包成带封面的正式报告(docx)

依赖:
  - pandoc (PATH 或本机常见路径)
  - python-docx
  - LibreOffice (仅 docx2pdf / package 可选)
"""
import argparse
import os
import shutil
import subprocess
import sys
import tempfile

HERE = os.path.dirname(os.path.abspath(__file__))
REF_DOC = os.path.join(HERE, "reference.docx")
SUBCOMMANDS = {"to-docx", "to-md", "gen-report", "batch", "docx2pdf", "package"}


# ---------------------------------------------------------------------------
# 环境探测
# ---------------------------------------------------------------------------
def find_pandoc():
    """定位 pandoc: 先查 PATH, 再查常见安装位置。"""
    found = shutil.which("pandoc")
    if found:
        return found
    candidates = [
        os.path.join(os.path.expanduser("~"), "bin", "pandoc.exe"),
        os.path.join(os.path.expanduser("~"), "bin", "pandoc"),
        r"C:\Users\honghuang\bin\pandoc.exe",
        r"C:\Program Files\Pandoc\pandoc.exe",
        "/usr/local/bin/pandoc",
        "/usr/bin/pandoc",
    ]
    for c in candidates:
        if c and os.path.isfile(c):
            return c
    return None


def find_soffice():
    """定位 LibreOffice soffice。"""
    for name in ("soffice", "libreoffice"):
        p = shutil.which(name)
        if p:
            return p
    candidates = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for c in candidates:
        if os.path.isfile(c):
            return c
    return None


def run_pandoc(args):
    pandoc = find_pandoc()
    if not pandoc:
        sys.stderr.write("[docx_utils] 未找到 pandoc, 请先安装并加入 PATH。\n")
        sys.exit(2)
    proc = subprocess.run([pandoc] + args, capture_output=True, text=True)
    if proc.returncode != 0:
        sys.stderr.write("[docx_utils] pandoc 执行失败:\n")
        sys.stderr.write(proc.stderr)
        sys.exit(1)
    return proc


def run_soffice(soffice, filepath, outdir):
    proc = subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", outdir, filepath],
        capture_output=True, text=True)
    if proc.returncode != 0:
        sys.stderr.write("[docx_utils] soffice 转换失败: %s\n" % filepath)
        sys.stderr.write(proc.stderr)
        sys.exit(1)


# ---------------------------------------------------------------------------
# 样式模板 / 页码 / 封面
# ---------------------------------------------------------------------------
def ensure_ref_doc():
    """首次运行生成中文友好样式模板 reference.docx。"""
    if os.path.exists(REF_DOC):
        return REF_DOC
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn

        doc = Document()
        normal = doc.styles["Normal"]
        normal.font.name = "Microsoft YaHei"
        normal.font.size = Pt(12)
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.append(rfonts)
        for a in ("w:ascii", "w:hAnsi", "w:eastAsia"):
            rfonts.set(qn(a), "Microsoft YaHei")

        levels = [
            ("Heading 1", 16, True, RGBColor(0x1F, 0x4E, 0x79)),
            ("Heading 2", 14, True, RGBColor(0x2E, 0x74, 0xB5)),
            ("Heading 3", 13, True, RGBColor(0x1F, 0x4E, 0x79)),
        ]
        for name, size, bold, color in levels:
            try:
                st = doc.styles[name]
                st.font.name = "Microsoft YaHei"
                st.font.size = Pt(size)
                st.font.bold = bold
                st.font.color.rgb = color
                er = st.element.get_or_add_rPr()
                rf = er.find(qn("w:rFonts"))
                if rf is None:
                    rf = er.makeelement(qn("w:rFonts"), {})
                    er.append(rf)
                for a in ("w:ascii", "w:hAnsi", "w:eastAsia"):
                    rf.set(qn(a), "Microsoft YaHei")
            except KeyError:
                pass
        for name in ("Title", "Author", "Date"):
            try:
                st = doc.styles[name]
                st.font.name = "Microsoft YaHei"
                st.font.size = Pt(18 if name == "Title" else 12)
                st.font.bold = (name == "Title")
                st.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
            except KeyError:
                pass
        doc.save(REF_DOC)
    except Exception as e:
        sys.stderr.write("[docx_utils] 生成样式模板失败(改用 pandoc 默认): %s\n" % e)
        return None
    return REF_DOC


def add_page_numbers(path):
    """给 docx 每个节页脚加 '第 X 页' 页码字段。"""
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document(path)

        def add_field(paragraph, field_code):
            run = paragraph.add_run()
            f1 = OxmlElement("w:fldChar")
            f1.set(qn("w:fldCharType"), "begin")
            it = OxmlElement("w:instrText")
            it.set(qn("xml:space"), "preserve")
            it.text = field_code
            f2 = OxmlElement("w:fldChar")
            f2.set(qn("w:fldCharType"), "end")
            run._r.append(f1)
            run._r.append(it)
            run._r.append(f2)

        for section in doc.sections:
            p = section.footer.paragraphs[0]
            p.text = ""
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("第 ")
            add_field(p, "PAGE")
            p.add_run(" 页")
        doc.save(path)
    except Exception as e:
        sys.stderr.write("[docx_utils] 添加页码失败(跳过): %s\n" % e)


def _run_props(size, bold, color_hex, font="Microsoft YaHei"):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        rFonts.set(qn(a), font)
    rPr.append(rFonts)
    if bold:
        rPr.append(OxmlElement("w:b"))
    color = OxmlElement("w:color")
    color.set(qn("w:val"), color_hex)
    rPr.append(color)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size))
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(size))
    rPr.append(szCs)
    return rPr


def _make_para(text, size, bold, color_hex, align="center", before=0, after=200):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(before))
    sp.set(qn("w:after"), str(after))
    pPr.append(sp)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), align)
    pPr.append(jc)
    p.append(pPr)
    r = OxmlElement("w:r")
    r.append(_run_props(size, bold, color_hex))
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def add_cover(path, title, author, date):
    """在 docx 开头插入封面页(标题/作者/日期 + 分页)。"""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document(path)
    body = doc.element.body
    title_p = _make_para(title or "报告", 36, True, "1F4E79", "center", 240, 200)
    author_p = _make_para(author or "", 24, False, "404040", "center", 0, 80)
    date_p = _make_para(date or "", 24, False, "404040", "center", 0, 200)
    pb = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    pb.append(r)
    for el in (pb, date_p, author_p, title_p):
        body.insert(0, el)
    doc.save(path)


# ---------------------------------------------------------------------------
# 核心功能 (CLI 与 GUI 共用)
# ---------------------------------------------------------------------------
def do_to_docx(md, out, toc=False, ref=None):
    ref = ref or ensure_ref_doc()
    pargs = [md, "-o", out]
    if ref:
        pargs += ["--reference-doc", ref]
    if toc:
        pargs += ["--toc"]
    run_pandoc(pargs)
    print("[docx_utils] 已生成: %s" % out)


def do_to_md(docx, out):
    run_pandoc([docx, "-o", out, "--wrap=none"])
    print("[docx_utils] 已生成: %s" % out)


def do_gen_report(md, out, title=None, author=None, date=None, ref=None, no_pagenum=False):
    ref = ref or ensure_ref_doc()
    pargs = [md, "-o", out, "--toc"]
    if ref:
        pargs += ["--reference-doc", ref]
    meta = []
    if title:
        meta.append(("title", title))
    if author:
        meta.append(("author", author))
    if date:
        meta.append(("date", date))
    for k, v in meta:
        pargs += ["-M", "%s=%s" % (k, v)]
    run_pandoc(pargs)
    if not no_pagenum:
        add_page_numbers(out)
    print("[docx_utils] 报告已生成: %s" % out)


def do_batch(directory, mode, recursive=False, out=None):
    if mode not in ("md2docx", "docx2md"):
        sys.stderr.write("[docx_utils] --mode 必须是 md2docx 或 docx2md\n")
        sys.exit(2)
    outdir = out or directory
    os.makedirs(outdir, exist_ok=True)
    if mode == "md2docx":
        src_ext, dst_ext = ".md", ".docx"
        ref = ensure_ref_doc()
    else:
        src_ext, dst_ext = ".docx", ".md"
        ref = None
    count = 0
    if recursive:
        walker = []
        for root, dirs, files in os.walk(directory):
            if os.path.basename(root).lower() == "archive":
                continue
            for f in files:
                walker.append(os.path.join(root, f))
    else:
        walker = [os.path.join(directory, f) for f in os.listdir(directory)]
    for src in walker:
        if not src.lower().endswith(src_ext):
            continue
        base = os.path.splitext(os.path.basename(src))[0]
        dst = os.path.join(outdir, base + dst_ext)
        if mode == "md2docx":
            pargs = [src, "-o", dst]
            if ref:
                pargs += ["--reference-doc", ref]
            run_pandoc(pargs)
        else:
            run_pandoc([src, "-o", dst, "--wrap=none"])
        count += 1
        print("[docx_utils] [%d] %s -> %s" % (count, src, dst))
    print("[docx_utils] 批量完成, 共 %d 个文件" % count)


def do_docx2pdf(src, out=None, install=False):
    soffice = find_soffice()
    if not soffice and install:
        print("[docx_utils] 尝试安装 LibreOffice ...")
        subprocess.run(["winget", "install", "--id", "TheDocumentFoundation.LibreOffice",
                        "--accept-package-agreements", "--accept-source-agreements", "--silent"],
                       check=False)
        soffice = find_soffice()
    if not soffice:
        sys.stderr.write("[docx_utils] 未找到 LibreOffice/soffice, 请先安装:\n"
                         "  winget install --id TheDocumentFoundation.LibreOffice --silent\n"
                         "安装后即可使用 docx -> PDF。\n")
        sys.exit(2)
    if os.path.isdir(src):
        outdir = out or src
        os.makedirs(outdir, exist_ok=True)
        files = [os.path.join(src, f) for f in os.listdir(src) if f.lower().endswith(".docx")]
        for f in files:
            run_soffice(soffice, f, outdir)
            print("[docx_utils] 已生成: %s.pdf" % os.path.splitext(f)[0])
    else:
        outdir = out or os.path.dirname(src)
        run_soffice(soffice, src, outdir)
        print("[docx_utils] 已生成: %s.pdf" % os.path.splitext(src)[0])


def gather_md_files(directory, recursive, order):
    if order:
        files = [os.path.join(directory, f.strip()) for f in order.split(",")]
        missing = [f for f in files if not os.path.isfile(f)]
        if missing:
            sys.stderr.write("[docx_utils] 以下文件不存在: %s\n" % ", ".join(missing))
            sys.exit(2)
        return files
    result = []
    if recursive:
        for root, dirs, files in os.walk(directory):
            if os.path.basename(root).lower() == "archive":
                continue
            for f in files:
                if f.lower().endswith(".md"):
                    result.append(os.path.join(root, f))
    else:
        for f in os.listdir(directory):
            if f.lower().endswith(".md"):
                result.append(os.path.join(directory, f))
    result.sort()
    return result


def do_package(directory, out, title=None, author=None, date=None,
               recursive=False, order=None, ref=None):
    files = gather_md_files(directory, recursive, order)
    if not files:
        sys.stderr.write("[docx_utils] 目录内未找到 .md 文件: %s\n" % directory)
        sys.exit(2)
    ref = ref or ensure_ref_doc()
    pargs = list(files) + ["-o", out, "--toc"]
    if ref:
        pargs += ["--reference-doc", ref]
    run_pandoc(pargs)
    if title or author or date:
        add_cover(out, title, author, date)
    add_page_numbers(out)
    print("[docx_utils] 已打包 %d 个文档 -> %s" % (len(files), out))


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------
def build_cli():
    p = argparse.ArgumentParser(prog="docx_utils.pyw",
                                description="docx 与 Markdown 互转 / 报告生成 (依赖 pandoc + python-docx)")
    sub = p.add_subparsers(dest="cmd", required=True)

    t = sub.add_parser("to-docx", help="Markdown -> docx")
    t.add_argument("input", help="输入 .md")
    t.add_argument("-o", "--out", required=True, help="输出 .docx")
    t.add_argument("--ref", help="自定义样式模板")
    t.add_argument("--toc", action="store_true", help="生成目录")
    t.set_defaults(func=lambda a: do_to_docx(a.input, a.out, a.toc, a.ref))

    m = sub.add_parser("to-md", help="docx -> Markdown")
    m.add_argument("input", help="输入 .docx")
    m.add_argument("-o", "--out", required=True, help="输出 .md")
    m.set_defaults(func=lambda a: do_to_md(a.input, a.out))

    r = sub.add_parser("gen-report", help="Markdown -> 正式报告 docx")
    r.add_argument("input", help="输入 .md")
    r.add_argument("-o", "--out", required=True, help="输出 .docx")
    r.add_argument("--title")
    r.add_argument("--author")
    r.add_argument("--date")
    r.add_argument("--ref")
    r.add_argument("--no-pagenum", action="store_true")
    r.set_defaults(func=lambda a: do_gen_report(a.input, a.out, a.title, a.author, a.date, a.ref, a.no_pagenum))

    b = sub.add_parser("batch", help="目录批量互转")
    b.add_argument("--dir", required=True, help="源目录")
    b.add_argument("--mode", required=True, choices=["md2docx", "docx2md"])
    b.add_argument("--recursive", action="store_true", help="递归子目录")
    b.add_argument("-o", "--out", help="输出目录(默认同输入目录)")
    b.set_defaults(func=lambda a: do_batch(a.dir, a.mode, a.recursive, a.out))

    d = sub.add_parser("docx2pdf", help="docx -> PDF (需 LibreOffice)")
    d.add_argument("input", help="输入 .docx 或目录")
    d.add_argument("-o", "--out", help="输出目录(默认同输入)")
    d.add_argument("--install", action="store_true", help="未安装时自动用 winget 安装 LibreOffice")
    d.set_defaults(func=lambda a: do_docx2pdf(a.input, a.out, a.install))

    pk = sub.add_parser("package", help="目录 Markdown 打包成带封面报告")
    pk.add_argument("--dir", required=True, help="源目录(含 .md)")
    pk.add_argument("-o", "--out", required=True, help="输出 .docx")
    pk.add_argument("--title")
    pk.add_argument("--author")
    pk.add_argument("--date")
    pk.add_argument("--recursive", action="store_true")
    pk.add_argument("--order", help="文件顺序, 逗号分隔(相对 --dir 的文件名)")
    pk.add_argument("--ref")
    pk.set_defaults(func=lambda a: do_package(a.dir, a.out, a.title, a.author, a.date,
                                              a.recursive, a.order, a.ref))
    return p


def cli_main():
    if not find_pandoc():
        sys.stderr.write("[docx_utils] 未找到 pandoc, 请先安装并加入 PATH。\n")
        sys.exit(2)
    args = build_cli().parse_args()
    args.func(args)


# ---------------------------------------------------------------------------
# GUI (双击 / 无参数)
# ---------------------------------------------------------------------------
def kill_previous():
    """单实例保护: 结束上一个本脚本的 GUI 实例。仅 Windows。"""
    if os.name != "nt":
        return
    lock = os.path.join(tempfile.gettempdir(), "docx_utils_gui.lock")
    prev = None
    if os.path.exists(lock):
        try:
            prev = int(open(lock, "r").read().strip())
        except Exception:
            prev = None
    if prev:
        r = subprocess.run(["tasklist", "/fi", "PID eq %d" % prev, "/nh"],
                           capture_output=True, text=True)
        if str(prev) in r.stdout:
            subprocess.run(["taskkill", "/pid", str(prev), "/f"], capture_output=True)
    try:
        open(lock, "w").write(str(os.getpid()))
    except Exception:
        pass


FOLDER_OPS = {"batch", "package"}
FILE_OPS = {"to-docx", "to-md", "gen-report"}


def gui_main():
    kill_previous()
    import tkinter as tk
    from tkinter import ttk, filedialog, messagebox

    root = tk.Tk()
    root.title("docx_utils - 文档转换工具")
    root.geometry("640x560")

    # 日志重定向
    class Redirect:
        def __init__(self, widget):
            self.widget = widget

        def write(self, s):
            self.widget.configure(state="normal")
            self.widget.insert("end", s)
            self.widget.see("end")
            self.widget.configure(state="disabled")

        def flush(self):
            pass

    op_var = tk.StringVar(value="to-docx")
    input_var = tk.StringVar()
    output_var = tk.StringVar()
    title_var = tk.StringVar()
    author_var = tk.StringVar()
    date_var = tk.StringVar()

    def browse_input():
        op = op_var.get()
        if op in FOLDER_OPS or (op == "docx2pdf" and False):
            p = filedialog.askdirectory(title="选择输入目录")
        else:
            p = filedialog.askopenfilename(
                title="选择输入文件",
                filetypes=[("Markdown", "*.md"), ("Word", "*.docx"), ("所有文件", "*.*")])
        if p:
            input_var.set(p)
            # 自动填写输出默认值
            if op in ("to-docx", "gen-report"):
                output_var.set(os.path.splitext(p)[0] + ".docx")
            elif op == "to-md":
                output_var.set(os.path.splitext(p)[0] + ".md")
            elif op == "docx2pdf":
                output_var.set(os.path.splitext(p)[0] + ".pdf")

    def browse_output():
        op = op_var.get()
        if op in FOLDER_OPS or op == "batch":
            p = filedialog.askdirectory(title="选择输出目录")
        else:
            p = filedialog.asksaveasfilename(
                title="选择输出文件",
                filetypes=[("Word", "*.docx"), ("Markdown", "*.md"),
                           ("PDF", "*.pdf"), ("所有文件", "*.*")])
        if p:
            output_var.set(p)

    def run_task():
        op = op_var.get()
        try:
            if op == "to-docx":
                do_to_docx(input_var.get(), output_var.get(),
                           toc=True, ref=None)
            elif op == "to-md":
                do_to_md(input_var.get(), output_var.get())
            elif op == "gen-report":
                do_gen_report(input_var.get(), output_var.get(),
                              title_var.get() or None, author_var.get() or None,
                              date_var.get() or None)
            elif op == "batch":
                mode = mode_var.get()
                do_batch(input_var.get(), mode, recursive_var.get(),
                         output_var.get() or None)
            elif op == "docx2pdf":
                do_docx2pdf(input_var.get(), output_var.get() or None)
            elif op == "package":
                do_package(input_var.get(), output_var.get(),
                           title_var.get() or None, author_var.get() or None,
                           date_var.get() or None, recursive_var.get())
            log.insert("end", "[完成]\n")
        except SystemExit as e:
            log.insert("end", "[退出码 %s]\n" % e.code)
        except Exception as e:
            log.insert("end", "[错误] %s\n" % e)

    # 布局
    tk.Label(root, text="操作:").grid(row=0, column=0, sticky="w", padx=8, pady=4)
    op_combo = ttk.Combobox(root, textvariable=op_var, state="readonly",
                            values=["to-docx", "to-md", "gen-report",
                                    "batch", "docx2pdf", "package"])
    op_combo.grid(row=0, column=1, columnspan=2, sticky="ew", padx=8, pady=4)

    tk.Label(root, text="输入:").grid(row=1, column=0, sticky="w", padx=8, pady=4)
    tk.Entry(root, textvariable=input_var).grid(row=1, column=1, sticky="ew", padx=8, pady=4)
    tk.Button(root, text="浏览", command=browse_input).grid(row=1, column=2, padx=8, pady=4)

    tk.Label(root, text="输出:").grid(row=2, column=0, sticky="w", padx=8, pady=4)
    tk.Entry(root, textvariable=output_var).grid(row=2, column=1, sticky="ew", padx=8, pady=4)
    tk.Button(root, text="浏览", command=browse_output).grid(row=2, column=2, padx=8, pady=4)

    # 报告/打包 专用选项
    tk.Label(root, text="标题:").grid(row=3, column=0, sticky="w", padx=8, pady=4)
    tk.Entry(root, textvariable=title_var).grid(row=3, column=1, columnspan=2, sticky="ew", padx=8, pady=4)
    tk.Label(root, text="作者:").grid(row=4, column=0, sticky="w", padx=8, pady=4)
    tk.Entry(root, textvariable=author_var).grid(row=4, column=1, columnspan=2, sticky="ew", padx=8, pady=4)
    tk.Label(root, text="日期:").grid(row=5, column=0, sticky="w", padx=8, pady=4)
    tk.Entry(root, textvariable=date_var).grid(row=5, column=1, columnspan=2, sticky="ew", padx=8, pady=4)

    # batch 模式选择
    mode_var = tk.StringVar(value="md2docx")
    recursive_var = tk.BooleanVar(value=False)
    mode_frame = tk.Frame(root)
    tk.Label(mode_frame, text="批量模式:").pack(side="left")
    ttk.Radiobutton(mode_frame, text="md->docx", variable=mode_var,
                    value="md2docx").pack(side="left", padx=4)
    ttk.Radiobutton(mode_frame, text="docx->md", variable=mode_var,
                    value="docx2md").pack(side="left", padx=4)
    tk.Checkbutton(mode_frame, text="递归子目录", variable=recursive_var).pack(side="left", padx=4)
    mode_frame.grid(row=6, column=0, columnspan=3, sticky="w", padx=8, pady=4)

    tk.Button(root, text="运行", command=run_task, bg="#2E74B5", fg="white",
              height=2).grid(row=7, column=0, columnspan=3, sticky="ew", padx=8, pady=8)

    log = tk.Text(root, height=14, state="disabled")
    log.grid(row=8, column=0, columnspan=3, sticky="nsew", padx=8, pady=4)
    scroll = tk.Scrollbar(root, command=log.yview)
    log.configure(yscrollcommand=scroll.set)
    scroll.grid(row=8, column=3, sticky="ns", pady=4)

    sys.stdout = Redirect(log)
    sys.stderr = Redirect(log)
    log.configure(state="normal")
    log.insert("end", "[docx_utils] 选择操作与文件后点'运行'。\n")
    log.configure(state="disabled")

    root.columnconfigure(1, weight=1)
    root.rowconfigure(8, weight=1)
    root.mainloop()


# ---------------------------------------------------------------------------
# 入口
# ---------------------------------------------------------------------------
def main():
    cli = "--cli" in sys.argv or (len(sys.argv) > 1 and sys.argv[1] in SUBCOMMANDS)
    if "--cli" in sys.argv:
        sys.argv.remove("--cli")
    if cli:
        cli_main()
    else:
        gui_main()


if __name__ == "__main__":
    main()
