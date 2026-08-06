#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
md_to_research_log_docx.pyw
===========================
将《E:/sEMG_B_Project/docs/研究日志.md》
转换为格式与《……研究日志_org.docx》保持一致的 Word 文档。

本脚本通过逐项解析 _org.docx 的真实 OOXML，把其版式规则固化下来，
以便后续更新 .md 后一键重新生成 .docx，且版面与原始 org 版本一致。

==== 已固化的版式规则 ====
1. 页面
   - A4（11900×16840 twips）；页边距 上860/右460/下1120/左740 twips，
     页眉距601、页脚距921 twips。
   - org 未引用任何 header/footer，因此生成文件也**不带页码**。

2. 封面页
   - 标题段用 `space_before=3080 twips` 把标题推到页面中上部（等效 org 由封面
     空段落下推的视觉位置），不再插入空段落。
   - 标题（`#` 文本 +「研究日志」）：微软正黑体、加粗、20.5pt、居中、
     行高 536 twips（精确）。
   - 封面信息行（`#` 与第一个 `##` 之间）：居中、宋体 9.5pt、加粗、颜色 2D2D2D；
     按字段类型分别复刻 org 的段前距与缩进（学校/人员/日期）。
   - 封面结束后插入**显式分页符**，正文从第 2 页开始。

3. 章节与子条目
   - `##`：应用「标题 1」样式（pStyle=1），并在首个 `##` 上加左缩进 1021、
     段前 53；后续 `##` 使用精确行高 376 twips。
   - `###`：仅加粗，去掉末尾句号。

4. 表格（核心）
   解析连续 `|` 行；空单元格按 org 约定视为「与左侧非空单元格合并」。
   已知表格类型匹配 org 的精确列宽/合并结构：
   - 类型 A（日志清单 时间/操作/内容/字数）：4 列，列宽 [2326,954,4553,967]。
   - 类型 B（横向环节分布）：12 物理列，6 个合并表头（每 2 列一组合并），
     合并单元格宽度 [1380,1440,1440,1440,1464,1464]。
   - 类型 C（时间分布：按周或按月）：首列「周」/「月份」，其余为分桶标签；列数随日期范围自动扩展，渲染时均匀分配列宽（不再固定 12 列）。
   - 类型 D（日志详情，表头 课题环节| |环节名| |，数据行 6 列）：
     表头两组合并（span=3），宽度 [4899,3735]；
     数据行常规 6 列宽度 [1729,1957,1213,1969,1225,541]；
     描述行（课题来源/期望效果/课题摘要…）第一列独占，第二列跨 5 列，
     宽度 [1729,6905]。
   - 其它表格：按内容长度比例分配宽度。
   所有表格统一：style=TableNormal、tblLayout=fixed、tblLook=01E0、
   tblCellMar 全 0、单元格四边 DBDEE4 0.75pt 边框、垂直居中、
   单元格段落段前按类型分行（log 127/13、dist 127/127/155、weekly 155、detail 127）、
   文字宋体 10.5pt、加粗（org 实测所有表格单元格均为 10.5pt 加粗）；
   水平对齐按各列规则（详见下方 ALIGN_* 常量）：日志内容列两端对齐、详情描述列左对齐、
   dist 数字行居中、其余居中，并非全表居中。
   ⚠️ 固定布局下「列宽由 tblGrid/gridCol 决定、非单元格 tcW」——每个表在
   apply_table_style 后必须 set_table_grid(table, widths) 覆盖默认等宽 gridCol，
   否则宽列被压成等宽、文字狂折行、整表变高溢出分页（曾导致多 ~5 页）。

5. 图片
   在「柱状图呈现」说明段后插入 assets/ 下的两张统计图，强制 15.24×6.0cm
   居中（与 org 一致）。图为静态快照，数据大幅变化时需手动更新。

交互：
  - .pyw 无控制台，双击启动 tkinter GUI 选择 .md，输出与 .md 同目录同名 .docx。
  - 也支持命令行：pythonw md_to_research_log_docx.pyw 输入.md [-o 输出.docx]
"""

import os
import re
import sys
import argparse
import traceback
import tempfile

from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- 路径默认值 ----------
DEFAULT_MD = r"E:\sEMG_B_Project\docs\研究日志.md"
HERE = os.path.dirname(os.path.abspath(__file__))

# ---------- 常量（单位：pt 或 twips） ----------
CJK = "宋体"
TITLE_FONT = "Microsoft JhengHei"
TITLE_PT = 20.5            # sz=41
COVER_PT = 9.5             # sz=19
BODY_PT = 10.5             # sz=21（org 实测：正文/表格/详情全为 10.5pt；但仅日志表数据行的 时间/操作/字数 三列加粗，其余均不加粗）
H1_PT = 13.0               # sz=26
GRID_COLOR = "DBDEE4"
GRID_SZ = 6                # 0.75pt
TEXT_COLOR = "2D2D2D"

TBL_WIDTH_DXA = 8800
TBL_IND_DXA = 1034
IMG_W_CM = 15.24
IMG_H_H_CM = 6.0

# 封面段参数（单位 twips）
TITLE_LINE = 536
COVER_SCHOOL_BEFORE = 53
COVER_PERSON_BEFORE = 147
COVER_PERSON_LINE = 609      # auto
COVER_DATE_BEFORE = 87
# 单元格段落段前距（twips），按 org 实测分类型/分行（org 基线值，必须一致）：
#  log    : 表头=127 / 数据行=13
#  dist   : 127, 127, 155（末行略大）
#  weekly : 全部=155
#  detail : 全部=127
CELL_SB_LOG_HEADER = 127
CELL_SB_LOG_DATA = 13
CELL_SB_DIST = [127, 127, 155]
CELL_SB_WEEKLY = 155
CELL_SB_DETAIL = 127
CELL_SPACING_BEFORE = 127  # 兜底（generic）

# 标题段前距（twips），用于把标题推到页面中上部（与 org 视觉位置对齐）
# 经 LibreOffice 渲染对比 org 封面标题顶边：gen 比 org 低 ~920 twips，
# 故取 3080（org 由 13 个空段落 + 标题自身间距推下，等效此处 space_before）。
TITLE_BEFORE = 3080

# 已知表格列宽（单位 twips）
WIDTHS_LOG = [1726, 954, 5153, 967]
# dist 表头合并宽度（每组 2 列合并），需与 WIDTHS_12GRID 对应两列之和一致
WIDTHS_DIST = [1467, 1467, 1467, 1467, 1466, 1466]
# weekly 数据行 12 个独立单元格宽度，与 WIDTHS_12GRID 一致
WIDTHS_WEEKLY = [734, 733, 734, 733, 734, 733, 734, 733, 733, 733, 733, 733]
WIDTHS_DETAIL_DATA = [1467, 1467, 1467, 1467, 1466, 1466]
WIDTHS_DETAIL_HEADER = [4401, 4399]  # 1467*3=4401, 1467+1466+1466=4399
# 12 物理列的 grid（dist 与 weekly 共用）：6 组 × 2 列，每组等宽 ≈1466.67 twips
# 组 0-3：734+733=1467，组 4-5：733+733=1466，总计 8800
WIDTHS_12GRID = [734, 733, 734, 733, 734, 733, 734, 733, 733, 733, 733, 733]

# 单元格水平对齐（org 实测，按物理列）。None 表示「不写 w:jc」，继承默认=左对齐，
# 与 org「未设 jc」的单元格 XML 等价。'both'=两端对齐(JUSTIFY)。
#   日志表：内容列(c2)两端对齐，其余居中
#   横向分布(dist)：全列左对齐（继承）；按周分布(weekly)：全列居中
#   详情表：所有列居中（直接读物理 XML 确认，含描述合并列）
ALIGN_LOG_H = ['center', 'center', 'both', 'center']          # 数据行：内容列两端对齐
ALIGN_LOG_V = ['center', 'center', None, 'center']
ALIGN_LOG_HEADER_H = ['center', 'center', 'center', 'center']  # 表头行全居中（含内容列）
ALIGN_LOG_HEADER_V = ['center', 'center', None, 'center']
ALIGN_WEEKLY_H = ['center'] * 12
ALIGN_WEEKLY_V = [None] * 12
ALIGN_DIST_H = [None] * 12
ALIGN_DIST_V = [None] * 12
ALIGN_DETAIL_H = ['center'] * 6
ALIGN_DETAIL_V = ['center', None, None, None, None, None]


# ---------- XML / 字体辅助 ----------
def _rPr(run):
    return run._element.get_or_add_rPr()


def set_cjk_font(run, font=CJK, size=None, bold=None, color=None):
    """设置 run 字体（含 eastAsia），并保证 w:rFonts 在 rPr 中的 schema 正确位置。"""
    rpr = _rPr(run)
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        idx = 0
        for i, ch in enumerate(rpr):
            if ch.tag == qn('w:rStyle'):
                idx = i + 1
        rpr.insert(idx, rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rfonts.set(qn(attr), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = None  # 先清再设
        color_el = rpr.find(qn('w:color'))
        if color_el is None:
            color_el = OxmlElement('w:color')
            rpr.append(color_el)
        color_el.set(qn('w:val'), color)


def set_paragraph_format(p, align=None, left_indent=None, right_indent=None,
                         first_line_indent=None, space_before=None,
                         space_after=None, line=None, line_rule=None):
    """统一走 python-docx 原生 paragraph_format，避免 XML 顺序错误。"""
    pf = p.paragraph_format
    if align is not None:
        p.alignment = align
    if left_indent is not None:
        pf.left_indent = Emu(left_indent * 635)
    if right_indent is not None:
        pf.right_indent = Emu(right_indent * 635)
    if first_line_indent is not None:
        pf.first_line_indent = Emu(first_line_indent * 635)
    if space_before is not None:
        pf.space_before = Emu(space_before * 635)
    if space_after is not None:
        pf.space_after = Emu(space_after * 635)
    if line is not None:
        if line_rule == 'exact':
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(line / 20.0)
        else:
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = line / 240.0
    else:
        # 显式清除行距，避免继承默认样式中的行距设置
        pPr = p._p.get_or_add_pPr()
        sp = pPr.find(qn('w:spacing'))
        if sp is not None:
            for a in ('w:line', 'w:lineRule'):
                if sp.get(qn(a)) is not None:
                    del sp.attrib[qn(a)]


def add_empty_para(doc):
    """添加一个完全空的段落（用于推标题位置）。"""
    return doc.add_paragraph()


def add_page_break(doc):
    """在段落内插入显式分页符。"""
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    return p


def neutralize_doc_defaults(doc):
    """覆写 python-docx 默认模板带来的 docDefaults，使其与 org 一致：
    - 段落段后距 after=0（默认 200 twips，会让每个单元格/段落下方多出空白，
      是所有表格变高、整文档多 4 页的根因）；
    - 不显式设置行距（line/lineRule 删除），沿用“单倍行距”（natural single），
      与 org（org 的 docDefaults 无 spacing 元素）完全一致。
    """
    styles_el = doc.styles.element
    dd = styles_el.find(qn('w:docDefaults'))
    if dd is None:
        dd = OxmlElement('w:docDefaults')
        styles_el.insert(0, dd)
    pPrDefault = dd.find(qn('w:pPrDefault'))
    if pPrDefault is None:
        pPrDefault = OxmlElement('w:pPrDefault')
        dd.append(pPrDefault)
    pPr = pPrDefault.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        pPrDefault.append(pPr)
    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        sp = OxmlElement('w:spacing')
        pPr.append(sp)
    sp.set(qn('w:after'), '0')
    for a in ('w:line', 'w:lineRule'):
        if sp.get(qn(a)) is not None:
            del sp.attrib[qn(a)]


def neutralize_heading1_style(doc):
    """把生成文档的『标题 1』样式对齐 org：段前 0、字号 13pt（sz=26）。
    否则非首个 ## 会继承默认 Heading 1 的 space_before=480 而下沉。"""
    try:
        st = doc.styles['Heading 1']
    except Exception:
        return
    el = st.element
    pPr = el.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        el.insert(0, pPr)
    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        sp = OxmlElement('w:spacing')
        pPr.insert(0, sp)
    sp.set(qn('w:before'), '0')
    sp.set(qn('w:after'), '0')
    rPr = el.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        el.append(rPr)
    sz = rPr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rPr.append(sz)
    sz.set(qn('w:val'), '26')


def is_md_only_text(text):
    """判断该行是否属于「当前 md 有、但 org 基线没有」的自由文本。

    org 正文除封面字段 / 章节标题 / 表格 / 图表注记外，不含其它自由文本。
    当前 md 额外出现的：每条日志详情前的【条目名】标号、汇总统计区的
    说明/按周分布/各阶段要点/周次列表等，org 中均不存在（已用 grep 核实
    计数为 0）。为与 org 版面一致，这些行在生成时直接丢弃。

    注意：md 中这些说明常以 Markdown 引用块 `> 说明：...` 形式出现，
    需先剥掉前导 `>` 引用符与 `**` 加粗包裹再做判断，否则会被误判为正常
    正文而泄漏进生成文档（实测曾导致多出 1 段说明文字）。
    """
    s = text.strip()
    s = re.sub(r'^>\s*', '', s).strip()   # 去掉 Markdown 引用符 >
    s = s.strip('*').strip()              # 去掉 ** 加粗包裹
    if re.match(r'^【.*】\s*$', s):
        return True
    if s.startswith('说明：') or s.startswith('说明:'):
        return True
    if '按周分布' in s:
        return True
    if s.startswith('各阶段要点'):
        return True
    if re.match(r'^-\s*第\d', s):
        return True
    return False


def add_markup(paragraph, text, size=BODY_PT, base_bold=False, color=None):
    """按 <br> 拆行；忽略 markdown 的 ** 加粗标记。

    org 的正文与表格中均不存在 markdown 加粗，加粗仅由显式规则控制：
    封面标题、二级目录(###)、日志表数据行的 时间/操作/字数 三列。因此这里
    直接把 ** 标记从文本中剥离（不渲染加粗），使显示文本与 org 完全一致。"""
    for seg_i, seg in enumerate(text.split('<br>')):
        if seg_i > 0:
            paragraph.add_run().add_break()
        clean = seg.replace('**', '')
        if clean == '':
            continue
        r = paragraph.add_run(clean)
        set_cjk_font(r, CJK, size, bold=base_bold, color=color)


# ---------- 表格辅助 ----------
def split_row(line):
    s = line.strip()
    if s.startswith('|'):
        s = s[1:]
    if s.endswith('|'):
        s = s[:-1]
    return [c.strip() for c in s.split('|')]


def is_separator(cells):
    return all(re.fullmatch(r':?-+:?', c) for c in cells) and len(cells) > 0


def merge_trailing_empties(row):
    """把行尾连续空单元格合并到左侧最后一个非空单元格（返回 [(text,span),...]）。"""
    out = []
    for c in row:
        if c == '':
            if out:
                out[-1] = (out[-1][0], out[-1][1] + 1)
            else:
                out.append(('', 1))
        else:
            out.append((c, 1))
    return out


def merge_header_empties(header, data_cols):
    """处理表头空单元格合并；确保合并后总列数 == data_cols。

    规则：空单元格向左合并；若表头合并后总列数 < data_cols，
    则把最后一个表头单元格的 span 扩展至填满 data_cols。
    """
    merged = []
    for c in header:
        if c == '':
            if merged:
                merged[-1] = (merged[-1][0], merged[-1][1] + 1)
            else:
                merged.append(('', 1))
        else:
            merged.append((c, 1))
    total = sum(s for _, s in merged)
    if total < data_cols and merged:
        last_text, last_span = merged[-1]
        merged[-1] = (last_text, last_span + (data_cols - total))
    return merged


def detect_table_type(header, data_rows):
    """根据表头与数据行列数判断表格类型。"""
    h = [c.strip() for c in header]
    max_cols = max(len(r) for r in data_rows) if data_rows else len(h)

    if h == ['时间', '操作', '内容', '字数']:
        return 'log'
    # 时间分布表（按周或按月）：首列固定为“周”或“月份”，其余为“第N周”或“YYYY年M月”
    # 标签（列数随日期范围自动扩展，不再硬编码 12 列）。只要首列是时间分桶标签且
    # 列数≥3（标签列 + 至少 1 个分桶 + 至少 1 个数据列）即判定为 weekly 表，
    # 渲染逻辑（均匀列宽、居中、合计行重算）对周/月通用。
    if h and h[0] in ('周', '月份') and len(h) >= 3:
        return 'weekly'
    # 横向环节分布：提出主题/文献调研/... 成对出现
    if (len(h) >= 10 and
        all(x in ['提出主题', '文献调研', '制定方案', '过程记录', '数据分析', '课题总结', ''] for x in h)):
        return 'dist'
    # 日志详情：表头 课题环节 | | 环节名 | |
    if (len(h) == 4 and h[0] == '课题环节' and h[1] == '' and h[3] == '' and
        max_cols == 6):
        return 'detail'
    return 'generic'


def set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tcPr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), str(GRID_SZ))
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), GRID_COLOR)
        borders.append(e)


def set_cell_margins(cell, top=0, left=0, bottom=0, right=0):
    """设置单元格级别的 tcMar（覆盖表格级 tblCellMar），单位 twips。"""
    tcPr = cell._tc.get_or_add_tcPr()
    mar = tcPr.find(qn('w:tcMar'))
    if mar is None:
        mar = OxmlElement('w:tcMar')
        tcPr.append(mar)
    for edge, val in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        e = mar.find(qn(f'w:{edge}'))
        if e is None:
            e = OxmlElement(f'w:{edge}')
            mar.append(e)
        e.set(qn('w:w'), str(val))
        e.set(qn('w:type'), 'dxa')


def set_gridspan(cell, span):
    """给单元格设置 gridSpan；span=1 时移除 gridSpan。"""
    tcPr = cell._tc.get_or_add_tcPr()
    gs = tcPr.find(qn('w:gridSpan'))
    if span <= 1:
        if gs is not None:
            tcPr.remove(gs)
        return
    if gs is None:
        gs = OxmlElement('w:gridSpan')
        tcPr.append(gs)
    gs.set(qn('w:val'), str(span))


def set_cell_width(cell, width_dxa):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(width_dxa))
    tcW.set(qn('w:type'), 'dxa')


def row_space_before(ttype, ri, nrows):
    """org 实测的单元格段落段前距（twips），按表格类型与行号返回。"""
    if ttype == 'log':
        return CELL_SB_LOG_HEADER if ri == 0 else CELL_SB_LOG_DATA
    if ttype == 'dist':
        if ri == 0:
            return 127
        if ri >= nrows - 1:
            return 155
        return 127
    if ttype == 'weekly':
        return CELL_SB_WEEKLY
    if ttype == 'detail':
        return CELL_SB_DETAIL
    return CELL_SPACING_BEFORE


def set_cell_formatting(cell, text, size=BODY_PT, bold=False, color=None,
                        space_before_tw=CELL_SPACING_BEFORE,
                        halign='center', valign='center'):
    """设置单元格：边框、垂直对齐、段落水平对齐、段前 space_before_tw、字体。

    halign: 'center' / 'both'(两端对齐) / 'left' / None(不写 w:jc，继承默认=左对齐)
    valign: 'center' / None(不写 w:vAlign，与 org「未设」单元格 XML 等价)
    二者默认均为 'center'（保持向后兼容）。
    """
    set_cell_border(cell)
    # 垂直对齐：None 表示不写，继承 org 默认（与 org XML 等价）
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn('w:vAlign'))
    if existing is not None:
        tcPr.remove(existing)
    if valign is not None:
        v = OxmlElement('w:vAlign')
        v.set(qn('w:val'), valign)
        tcPr.append(v)
    # 段落格式
    align_map = {'center': WD_ALIGN_PARAGRAPH.CENTER,
                 'both': WD_ALIGN_PARAGRAPH.JUSTIFY,
                 'left': WD_ALIGN_PARAGRAPH.LEFT}
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Emu(space_before_tw * 635)
        if halign is not None:
            p.alignment = align_map.get(halign, WD_ALIGN_PARAGRAPH.CENTER)
        # halign is None -> 不设置，继承默认（左对齐）
    add_markup(cell.paragraphs[0], text, size=size, base_bold=bold, color=color)
    # 空单元格兜底：org 的空单元格仍带 宋体/z/sz/b 的 run（仅无文字），
    # 若 add_markup 因 text 为空而未产生 run，这里补一个带格式的空 run，
    # 保证空单元格的 rPr 与 org 一致（避免 NO_RPR 差异）。
    if not cell.paragraphs[0].runs:
        er = cell.paragraphs[0].add_run('')
        set_cjk_font(er, CJK, size, bold=bold, color=color)


def apply_table_style(table):
    """设置 org 的表格级属性：style、layout、look、width、indent、cell margin=0。"""
    tblPr = table._tbl.tblPr

    # style = TableNormal
    style_el = tblPr.find(qn('w:tblStyle'))
    if style_el is None:
        style_el = OxmlElement('w:tblStyle')
        tblPr.insert(0, style_el)
    style_el.set(qn('w:val'), 'TableNormal')

    # tblW
    w = tblPr.find(qn('w:tblW'))
    if w is None:
        w = OxmlElement('w:tblW')
        tblPr.append(w)
    w.set(qn('w:w'), str(TBL_WIDTH_DXA))
    w.set(qn('w:type'), 'dxa')

    # jc
    jc = tblPr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc')
        tblPr.append(jc)
    jc.set(qn('w:val'), 'center')

    # tblInd
    ind = tblPr.find(qn('w:tblInd'))
    if ind is None:
        ind = OxmlElement('w:tblInd')
        tblPr.append(ind)
    ind.set(qn('w:w'), str(TBL_IND_DXA))
    ind.set(qn('w:type'), 'dxa')

    # tblLayout = fixed
    layout = tblPr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tblPr.append(layout)
    layout.set(qn('w:type'), 'fixed')

    # tblLook
    look = tblPr.find(qn('w:tblLook'))
    if look is None:
        look = OxmlElement('w:tblLook')
        tblPr.append(look)
    look.set(qn('w:val'), '01E0')
    look.set(qn('w:firstRow'), '1')
    look.set(qn('w:lastRow'), '1')
    look.set(qn('w:firstColumn'), '1')
    look.set(qn('w:lastColumn'), '1')
    look.set(qn('w:noHBand'), '0')
    look.set(qn('w:noVBand'), '0')

    # tblCellMar = 0（这是表格紧凑的关键）
    mar = tblPr.find(qn('w:tblCellMar'))
    if mar is None:
        mar = OxmlElement('w:tblCellMar')
        tblPr.append(mar)
    for edge in ('top', 'left', 'bottom', 'right'):
        e = mar.find(qn(f'w:{edge}'))
        if e is None:
            e = OxmlElement(f'w:{edge}')
            mar.append(e)
        e.set(qn('w:w'), '0')
        e.set(qn('w:type'), 'dxa')


def set_table_grid(table, widths):
    """固定布局(tblLayout=fixed)下，列宽由 tblGrid/gridCol 决定，
    LibreOffice/Word 不会用单元格 tcW，而会用 gridCol。必须把 gridCol
    宽度设为与列宽一致，否则表格会被渲染成 python-docx 默认的等宽列，
    导致宽列（如日志“内容”列）过窄、文字狂折行、整表变高溢出分页。"""
    tbl = table._tbl
    grid = tbl.find(qn('w:tblGrid'))
    if grid is None:
        grid = OxmlElement('w:tblGrid')
        tbl.insert(0, grid)
    for gc in grid.findall(qn('w:gridCol')):
        grid.remove(gc)
    for w in widths:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        gc.set(qn('w:type'), 'dxa')
        grid.append(gc)


def even_widths(ncol, total=TBL_WIDTH_DXA):
    """把 total 宽度（twips）均分给 ncol 列（固定布局下驱动列宽的 gridCol），
    末列修正舍入误差，保证各列之和 == total。用于「按周分布」等列数不固定的表。

    同时既作为 tblGrid/gridCol 宽度，也作为单元格 tcW 宽度，二者一致，
    避免 LibreOffice/Word 用 gridCol 渲染而 tcW 不一致导致的列宽漂移。
    """
    if ncol <= 0:
        return []
    base = total // ncol
    widths = [base] * ncol
    diff = total - sum(widths)
    if diff:
        widths[-1] += diff
    return widths


def width_by_content(header, data_rows, total=TBL_WIDTH_DXA):
    """按每列最大文字长度比例分配列宽（兜底）。"""
    n = max(len(header), max((len(r) for r in data_rows), default=len(header)))
    lens = [0] * n
    for i, c in enumerate(header):
        lens[i] = max(lens[i], len(c))
    for r in data_rows:
        for i, c in enumerate(r):
            lens[i] = max(lens[i], len(c))
    # 加底数避免太窄
    weights = [max(2, x) for x in lens]
    s = sum(weights)
    widths = [int(total * w / s) for w in weights]
    # 修正舍入误差
    diff = total - sum(widths)
    if widths and diff:
        widths[-1] += diff
    return widths


def _cell_to_int(s):
    """表格单元格文本转整数；非数字（如 —、空、占位符）记为 0。"""
    s2 = (s or '').replace('*', '').strip()
    if s2.isdigit():
        return int(s2)
    return 0


def build_table(doc, header, data_rows):
    """根据表头/数据构造与 org 版式一致的表格。"""
    ttype = detect_table_type(header, data_rows)

    if ttype == 'log':
        ncol = 4
        widths = WIDTHS_LOG
        nrows = 1 + len(data_rows)
        table = doc.add_table(rows=nrows, cols=ncol)
        apply_table_style(table)
        set_table_grid(table, WIDTHS_LOG)
        # 内容列(c=2)左右边距，避免文字贴边框
        CONTENT_COL_MARGIN = 80  # twips ≈ 5pt
        # 表头
        sb0 = row_space_before('log', 0, nrows)
        for c in range(ncol):
            cell = table.rows[0].cells[c]
            set_cell_width(cell, widths[c])
            if c == 2:
                set_cell_margins(cell, left=CONTENT_COL_MARGIN, right=CONTENT_COL_MARGIN)
            set_cell_formatting(cell, header[c], space_before_tw=sb0,
                                halign=ALIGN_LOG_HEADER_H[c], valign=ALIGN_LOG_HEADER_V[c],
                                bold=False)
        # 数据
        for ri, row in enumerate(data_rows):
            sb = row_space_before('log', ri + 1, nrows)
            for c in range(ncol):
                cell = table.rows[1 + ri].cells[c]
                set_cell_width(cell, widths[c])
                if c == 2:
                    set_cell_margins(cell, left=CONTENT_COL_MARGIN, right=CONTENT_COL_MARGIN)
                # 用户偏好：日志表数据行整表不加粗（含 时间/操作/字数），
                # 仅保留对齐规则（c0/c1/c3 居中、内容列 c2 两端对齐）。
                # 注：此点与 _org.docx 原样式不同（org 原对这 3 列加粗）。
                set_cell_formatting(cell, row[c] if c < len(row) else '',
                                    space_before_tw=sb,
                                    halign=ALIGN_LOG_H[c], valign=ALIGN_LOG_V[c],
                                    bold=False)
        return table

    if ttype == 'weekly':
        # 列数随日期范围自动扩展（周 + N 个“第N周”数据列），不再硬编码 12 列。
        ncol = len(header)
        widths = even_widths(ncol)
        nrows = 1 + len(data_rows)
        table = doc.add_table(rows=nrows, cols=ncol)
        apply_table_style(table)
        set_table_grid(table, widths)
        sb0 = row_space_before('weekly', 0, nrows)
        for c in range(ncol):
            cell = table.rows[0].cells[c]
            set_cell_width(cell, widths[c])
            set_cell_formatting(cell, header[c], space_before_tw=sb0,
                                halign='center', valign=None,
                                bold=False)
        # 重算“合计”行：合计 = 新增 + 修订（按列），依据实际条数更新；
        # 覆盖 md 中人工写死的“合计”行，避免与新增/修订不一致。
        try:
            _labels = [c.replace('*', '').strip() for c in (r[0] for r in data_rows)]
            _ai = _labels.index('新增')
            _ri = _labels.index('修订')
            _ti = _labels.index('合计')
        except (ValueError, IndexError):
            pass
        else:
            _new = [data_rows[_ti][0]]  # 保留原标签（可能含 **合计**）
            for c in range(1, ncol):
                _a = _cell_to_int(data_rows[_ai][c])
                _r = _cell_to_int(data_rows[_ri][c])
                _new.append(str(_a + _r))
            data_rows[_ti] = _new
        for ri, row in enumerate(data_rows):
            sb = row_space_before('weekly', ri + 1, nrows)
            for c in range(ncol):
                cell = table.rows[1 + ri].cells[c]
                set_cell_width(cell, widths[c])
                set_cell_formatting(cell, row[c] if c < len(row) else '',
                                    space_before_tw=sb,
                                    halign='center', valign=None,
                                    bold=False)
        return table

    if ttype == 'dist':
        # 12 物理列，表头 6 个合并单元格（每格 span=2）
        ncol = 12
        nrows = 1 + len(data_rows)
        table = doc.add_table(rows=nrows, cols=ncol)
        apply_table_style(table)
        set_table_grid(table, WIDTHS_12GRID)
        merged_header = merge_header_empties(header, ncol)
        col = 0
        sb0 = row_space_before('dist', 0, nrows)
        for vi, (text, span) in enumerate(merged_header):
            if span > 1:
                cell = table.rows[0].cells[col].merge(table.rows[0].cells[col + span - 1])
            else:
                cell = table.rows[0].cells[col]
            set_cell_width(cell, WIDTHS_DIST[vi])
            # 用户偏好：次数统计表整表居中（含表头与“新增/修订”行）。
            # 注：此点与 _org.docx 原样式不同（org 原表头/标签行左对齐、仅末行数字居中）。
            set_cell_formatting(cell, text, space_before_tw=sb0,
                                halign='center', valign=ALIGN_DIST_V[col],
                                bold=False)
            col += span
        # 数据行（12 个独立单元格）：整表居中
        for ri, row in enumerate(data_rows):
            sb = row_space_before('dist', ri + 1, nrows)
            for c in range(ncol):
                cell = table.rows[1 + ri].cells[c]
                set_cell_width(cell, WIDTHS_WEEKLY[c])
                set_cell_formatting(cell, row[c] if c < len(row) else '',
                                    space_before_tw=sb,
                                    halign='center', valign=ALIGN_DIST_V[c],
                                    bold=False)
        return table

    if ttype == 'detail':
        # 6 物理列，表头 2 组合并（span=3）
        ncol = 6
        nrows = 1 + len(data_rows)
        table = doc.add_table(rows=nrows, cols=ncol)
        apply_table_style(table)
        set_table_grid(table, WIDTHS_DETAIL_DATA)
        # 表头：取第二个非空标签
        section_name = header[2] if len(header) > 2 and header[2] else header[0]
        header_cells = [('课题环节', 3), (section_name, 3)]
        col = 0
        sb0 = row_space_before('detail', 0, nrows)
        for vi, (text, span) in enumerate(header_cells):
            if span > 1:
                cell = table.rows[0].cells[col].merge(table.rows[0].cells[col + span - 1])
            else:
                cell = table.rows[0].cells[col]
            set_cell_width(cell, WIDTHS_DETAIL_HEADER[vi])
            set_cell_formatting(cell, text, space_before_tw=sb0,
                                halign=ALIGN_DETAIL_H[col], valign=ALIGN_DETAIL_V[col],
                                bold=False)
            col += span
        # 数据行
        DETAIL_CONTENT_MARGIN = 80  # twips ≈ 5pt，描述列左右边距
        for ri, row in enumerate(data_rows):
            sb = row_space_before('detail', ri + 1, nrows)
            merged = merge_trailing_empties(row)
            # 保证总 span = ncol
            total_span = sum(s for _, s in merged)
            if total_span < ncol and merged:
                t, s = merged[-1]
                merged[-1] = (t, s + (ncol - total_span))
            col = 0
            idx = 0
            for text, span in merged:
                # 为被合并的单元格定位到起始单元格
                cell = table.rows[1 + ri].cells[col]
                if span > 1:
                    end_cell = table.rows[1 + ri].cells[col + span - 1]
                    cell.merge(end_cell)
                # 按起始列给宽度
                if col == 0:
                    set_cell_width(cell, WIDTHS_DETAIL_DATA[0])
                elif span == 5:
                    desc_width = sum(WIDTHS_DETAIL_DATA[1:])
                    set_cell_width(cell, desc_width)
                else:
                    set_cell_width(cell, WIDTHS_DETAIL_DATA[col])
                # 描述列(span=5)加左右内边距，避免文字贴边框
                if span == 5:
                    set_cell_margins(cell, left=DETAIL_CONTENT_MARGIN, right=DETAIL_CONTENT_MARGIN)
                # org 规则：标签列(c0)与短值列(span=1)居中；描述列(span=5，跨 5 列
                # 的长文本)左对齐（不写 w:jc，继承默认）。
                halign = None if span == 5 else 'center'
                set_cell_formatting(cell, text, space_before_tw=sb,
                                    halign=halign, valign=ALIGN_DETAIL_V[col],
                                    bold=False)
                col += span
                idx += 1
        return table

    # generic fallback
    ncol = max(len(header), max((len(r) for r in data_rows), default=len(header)))
    nrows = 1 + len(data_rows)
    widths = width_by_content(header, data_rows)
    table = doc.add_table(rows=nrows, cols=ncol)
    apply_table_style(table)
    set_table_grid(table, widths)
    sb0 = row_space_before('generic', 0, nrows)
    for c in range(ncol):
        cell = table.rows[0].cells[c]
        set_cell_width(cell, widths[c])
        set_cell_formatting(cell, header[c] if c < len(header) else '',
                            space_before_tw=sb0, bold=False)
    for ri, row in enumerate(data_rows):
        sb = row_space_before('generic', ri + 1, nrows)
        for c in range(ncol):
            cell = table.rows[1 + ri].cells[c]
            set_cell_width(cell, widths[c])
            set_cell_formatting(cell, row[c] if c < len(row) else '',
                                space_before_tw=sb, bold=False)
    return table


# ---------- 封面辅助 ----------
def add_cover_field(doc, text, field_type='date', before_override=None, line_override=None):
    """按 org 的字段类型插入封面信息行。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 为避免长占位文字（如"（填写作者姓名）"）在窄缩进内折行，
    # 所有封面字段统一使用较宽缩进；视觉仍居中。
    if field_type == 'school':
        before = COVER_SCHOOL_BEFORE
        indent_l, indent_r = 2613, 2612
        line = None
    elif field_type == 'person':
        before = COVER_PERSON_BEFORE
        indent_l, indent_r = 2619, 2612
        line = None
    else:  # date / 其它
        before = COVER_DATE_BEFORE
        indent_l, indent_r = 2619, 2612
        line = None
    if before_override is not None:
        before = before_override
    if line_override is not None:
        line = line_override

    kwargs = dict(space_before=before, left_indent=indent_l, right_indent=indent_r)
    if line is not None:
        kwargs['line'] = line   # 默认 MULTIPLE（org 无 lineRule = auto）
    set_paragraph_format(p, **kwargs)

    r = p.add_run(text)
    set_cjk_font(r, CJK, COVER_PT, bold=False, color=TEXT_COLOR)


# ---------- 图片 ----------
def add_image_centered(doc, path):
    if not os.path.exists(path):
        sys.stderr.write(f"[WARN] 图表缺失，跳过：{path}\n")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Cm(IMG_W_CM), height=Cm(IMG_H_H_CM))


# ---------- 时间分布柱状图（按周或按月，内联数据驱动，无需外部图片资源） ----------
def parse_weekly_table(md_path):
    """从 md 的「所有日志时间分布图」表格解析时间分布数据，返回 (labels, added, revised)。
    支持按周（表头含“第N周”）与按月（表头含“YYYY年M月”）两种表头。
    解析失败返回 (None, None, None)。图表完全由 md 表格数据驱动，不再依赖冻结的 PNG。"""
    try:
        with open(md_path, 'r', encoding='utf-8') as f:
            lines = f.read().split('\n')
    except OSError:
        return None, None, None
    hdr = None
    for idx, ln in enumerate(lines):
        if not ln.strip().startswith('|'):
            continue
        hc = [c.strip() for c in ln.split('|')]
        if hc and hc[0] == '':
            hc = hc[1:]
        if hc and hc[-1] == '':
            hc = hc[:-1]
        # 时间分布表：首列须为「周」或「月份」，且除首列外至少 2 个时间标签
        if not hc or hc[0] not in ('周', '月份'):
            continue
        cnt = sum(1 for c in hc[1:] if re.search(r'第\d+周|\d{4}年\d+月', c))
        if cnt >= 2:
            hdr = idx
            break
    if hdr is None:
        return None, None, None
    hcells = [c.strip() for c in lines[hdr].split('|')]
    if hcells and hcells[0] == '':
        hcells = hcells[1:]
    if hcells and hcells[-1] == '':
        hcells = hcells[:-1]
    weeks = []
    for c in hcells[1:]:  # 跳过首列「周」/「月份」
        m = re.search(r'第\d+周|\d{4}年\d+月', c)
        if m:
            weeks.append(m.group(0))
    if len(weeks) < 2:
        return None, None, None

    def parse_nums(ln):
        cs = [c.strip() for c in ln.split('|')]
        if cs and cs[0] == '':
            cs = cs[1:]
        if cs and cs[-1] == '':
            cs = cs[:-1]
        out = []
        for c in cs[1:]:
            c2 = c.replace('*', '').strip()
            if c2.isdigit():
                out.append(int(c2))
        return out

    # 表头与数据行之间可能有 markdown 分隔行（|---|），故显式定位「新增」行
    add_idx = None
    for j in range(hdr + 1, min(hdr + 5, len(lines))):
        if '新增' in lines[j]:
            add_idx = j
            break
    if add_idx is None:
        return None, None, None
    added = parse_nums(lines[add_idx])
    revised = parse_nums(lines[add_idx + 1])
    if len(added) != len(weeks) or len(revised) != len(weeks):
        return None, None, None
    return weeks, added, revised


def build_weekly_chart_png(md_path):
    """解析 md 时间分布表（按周或按月），生成中文柱状图到临时 PNG，返回路径
    （解析/生成失败返回 None）。"""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from matplotlib import font_manager as fm

    weeks, added, revised = parse_weekly_table(md_path)
    if not weeks:
        return None

    is_monthly = any('年' in w for w in weeks)

    FONT = r"C:\Users\honghuang\AppData\Local\Microsoft\Windows\Fonts\SarasaFixedSC-Regular.ttf"
    if os.path.exists(FONT):
        fm.fontManager.addfont(FONT)
        zh = fm.FontProperties(fname=FONT)
    else:
        zh = fm.FontProperties(family=['SimHei', 'Microsoft YaHei', 'STSong', 'sans-serif'])
    plt.rcParams['axes.unicode_minus'] = False

    total = sum(added) + sum(revised)
    x = range(len(weeks))
    w = 0.38
    fig, ax = plt.subplots(figsize=(6.0, 2.362), dpi=220)
    b1 = ax.bar([i - w / 2 for i in x], added, width=w, color="#2E75B6", label="新增日志")
    b2 = ax.bar([i + w / 2 for i in x], revised, width=w, color="#ED7D31", label="修订日志")
    ax.set_xticks(list(x))
    # 按月标签（如 2026年2月）较长，字体略小；按周标签（第N周）用原字号
    xtick_fs = 6 if is_monthly else 7
    ax.set_xticklabels(weeks, fontproperties=zh, fontsize=xtick_fs)
    ax.set_ylabel("条数", fontproperties=zh, fontsize=8)
    unit = '按月' if is_monthly else '按周'
    ax.set_title(f"日志时间分布图（{unit}，合计 {total} 条）", fontproperties=zh, fontsize=10)
    ax.legend(prop=zh, fontsize=7, loc="upper left")
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.tick_params(axis='y', labelsize=7)
    # 按月时单月条数可能远多于按周，ylim 动态取数据最大值（至少留 2 余量）
    ax.set_ylim(0, max(max(added), max(revised)) + 2)
    for bars in (b1, b2):
        for r in bars:
            h = r.get_height()
            if h > 0:
                ax.annotate(str(int(h)), (r.get_x() + r.get_width() / 2, h),
                            ha='center', va='bottom', fontsize=6, fontproperties=zh)
    fig.tight_layout()
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.png', prefix='wkchart_')
    tmp.close()
    fig.savefig(tmp.name, dpi=220)
    plt.close(fig)
    return tmp.name


# ---------- 主构建流程 ----------
def parse_and_build(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.read().split('\n')

    doc = Document()

    # 覆写默认 docDefaults / 标题 1 样式，使其与 org 一致（节行距、段后距等）
    neutralize_doc_defaults(doc)
    neutralize_heading1_style(doc)

    # 默认 Normal 样式：宋体 11pt
    normal = doc.styles['Normal']
    normal.font.size = Pt(BODY_PT)
    rfonts = normal.font._element.get_or_add_rPr().find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        normal.font._element.get_or_add_rPr().insert(0, rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        rfonts.set(qn(attr), CJK)

    # 页面设置
    sec = doc.sections[0]
    sec.page_width = Emu(11900 * 635)
    sec.page_height = Emu(16840 * 635)
    sec.top_margin = Emu(860 * 635)
    sec.bottom_margin = Emu(1120 * 635)
    sec.left_margin = Emu(740 * 635)
    sec.right_margin = Emu(460 * 635)
    sec.header_distance = Emu(601 * 635)
    sec.footer_distance = Emu(921 * 635)

    title_emitted = False
    cover_active = True
    inserted_images = False
    cover_date_count = 0   # 封面 date 字段出现次序（用于复刻 org 的空行/段前距）
    h1_count = 0
    in_detail_section = False   # 是否处于“日志详情”分页区
    last_was_heading = False    # 仅在详情分页区内用于判断是否需要在块前插分页符

    # 周分布实际总条数（新增+修订之和）：用于覆盖正文“合计 N 条”与表格“合计”行，
    # 使统计始终依据 md 中的新增/修订实际条数，而非人工写死的合计。
    _w, _a, _r = parse_weekly_table(md_path)
    weekly_total = (sum(_a) + sum(_r)) if _w is not None else None

    def maybe_insert_chart(text):
        nonlocal inserted_images
        if (not inserted_images) and ('柱状图呈现' in text or '时间分布图' in text):
            # 数据驱动：从 md 周分布表实时生成中文柱状图，不再依赖 assets 下的冻结 PNG
            png = build_weekly_chart_png(md_path)
            if png:
                add_image_centered(doc, png)
                try:
                    os.remove(png)
                except OSError:
                    pass
            inserted_images = True

    i = 0
    n = len(lines)
    while i < n:
        raw = lines[i]
        stripped = raw.strip()

        # 一级标题（首个 # ）
        if stripped.startswith('# ') and not stripped.startswith('##') and not title_emitted:
            main_title = stripped[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_format(p, left_indent=500, right_indent=500,
                                 space_before=TITLE_BEFORE,
                                 line=TITLE_LINE, line_rule='exact')
            # 第一行：主标题
            r1 = p.add_run(main_title)
            set_cjk_font(r1, CJK, TITLE_PT, bold=True)
            # 强制换行
            p.add_run().add_break()
            # 第二行：研究日志
            r2 = p.add_run('研究日志')
            set_cjk_font(r2, CJK, TITLE_PT, bold=True)
            title_emitted = True
            cover_active = True
            i += 1
            continue

        # 章节标题 ##
        if stripped.startswith('## '):
            cover_active = False
            # 每个 ## 前插入分页符（含首个：封面→正文），与 org 一致。
            add_page_break(doc)
            title_text = stripped[3:].strip()
            # 进入/退出“日志详情”分页区（该区每个详情表独占一页）。
            in_detail_section = ('详情' in title_text)
            last_was_heading = True

            p = doc.add_paragraph()
            # 应用「标题 1」样式
            try:
                p.style = doc.styles['Heading 1']
            except Exception:
                pass
            if h1_count == 0:
                set_paragraph_format(p, left_indent=1021, first_line_indent=0,
                                     space_before=53)
            else:
                set_paragraph_format(p, left_indent=1021, first_line_indent=0,
                                     space_before=0, line=376, line_rule='exact')
            r = p.add_run(title_text)
            set_cjk_font(r, CJK, H1_PT, bold=True, color=TEXT_COLOR)
            h1_count += 1
            i += 1
            continue

        # 子条目 ###
        if stripped.startswith('### '):
            cover_active = False
            txt = stripped[4:].strip().rstrip('。')
            # 详情分页区：首个 ###（紧跟本区 ##）不插分页符，其余每个 ### 前分页。
            if in_detail_section and not last_was_heading:
                add_page_break(doc)
            p = doc.add_paragraph()
            # org 的二级目录（###）通过 w:outlineLvl=1 进入导航窗格（非 Heading 样式）。
            pPr = p._p.get_or_add_pPr()
            for ex in pPr.findall(qn('w:outlineLvl')):
                pPr.remove(ex)
            ol = OxmlElement('w:outlineLvl')
            ol.set(qn('w:val'), '1')
            pPr.append(ol)
            r = p.add_run(txt)
            set_cjk_font(r, CJK, BODY_PT, bold=True)
            if in_detail_section:
                last_was_heading = True
            i += 1
            continue

        # 表格
        if stripped.startswith('|'):
            table_lines = []
            while i < n and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            if len(table_lines) < 2:
                for tl in table_lines:
                    p = doc.add_paragraph()
                    add_markup(p, tl, size=BODY_PT)
                continue
            rows = []
            header = None
            for tl in table_lines:
                cells = split_row(tl)
                if is_separator(cells):
                    continue
                if header is None:
                    header = cells
                else:
                    rows.append(cells)
            if header and rows:
                ttype = detect_table_type(header, rows)
                # 详情分页区：每个详情表独占一页；首个详情表（紧跟 ###）不插分页符。
                if ttype == 'detail' and in_detail_section and not last_was_heading:
                    add_page_break(doc)
                build_table(doc, header, rows)
                last_was_heading = False
            continue

        # 水平分隔线 --- （丢弃）
        if re.fullmatch(r'-{3,}', stripped):
            i += 1
            continue

        # 空行
        if stripped == '':
            i += 1
            continue

        # 封面块
        if cover_active:
            txt = re.sub(r'\*\*', '', stripped)
            # 判断字段类型
            if txt.startswith('学校'):
                # 标题→学校间用空行(段前10)撑开，与后续信息行间隔模式一致
                add_empty_para(doc).paragraph_format.space_before = Emu(10 * 635)
                add_cover_field(doc, txt, 'school', before_override=0)
                add_empty_para(doc).paragraph_format.space_before = Emu(10 * 635)
            elif txt.startswith(('作者', '指导老师')):
                # 段前 0，后接空行(段前10)，与最后3行间隔模式一致
                add_cover_field(doc, txt, 'person', before_override=0)
                add_empty_para(doc).paragraph_format.space_before = Emu(10 * 635)
            else:
                # date 字段：统一段前距与空行结构，六行信息间隔相等
                cover_date_count += 1
                if cover_date_count == 1:
                    add_cover_field(doc, txt, 'date', before_override=0)
                    add_empty_para(doc).paragraph_format.space_before = Emu(10 * 635)
                elif cover_date_count == 2:
                    add_cover_field(doc, txt, 'date', before_override=0)
                    add_empty_para(doc).paragraph_format.space_before = Emu(10 * 635)
                else:
                    add_cover_field(doc, txt, 'date', before_override=0)
            i += 1
            continue

        # 普通正文段落
        cover_active = False
        # 丢弃 org 基线中不存在的自由文本（【条目名】/说明/按周分布/各阶段要点/周次列表）
        if is_md_only_text(stripped):
            i += 1
            continue
        clean = stripped.strip('*').strip()
        # 覆盖正文中的“合计 N 条”为实际总条数（新增+修订之和），依据实际条数更新
        if weekly_total is not None:
            clean = re.sub(r'合计\s*\d+\s*条', f'合计 {weekly_total} 条', clean)
        p = doc.add_paragraph()
        add_markup(p, clean, size=BODY_PT)
        maybe_insert_chart(clean)
        i += 1

    normalize_empty_cells(doc)
    doc.save(docx_path)
    return docx_path


def normalize_empty_cells(doc):
    """org 中合并产生的空白续格（续列 tc）仍带一个 宋体/10.5pt/加粗 的空 run；
    python-docx 的 merge / 空单元格可能产生「无 run」或「run 无 rPr」的单元格。
    为与 org 字节级一致，遍历所有物理 tc：给无 run 的单元格补一个带格式空 run，
    给已有 run 但缺 rPr 的（多为空 run）补上 rPr。仅影响空白续格，不改动任何
    有内容的单元格。"""
    for t in doc.tables:
        tbl = t._tbl
        for tr in tbl.findall(qn('w:tr')):
            for tc in tr.findall(qn('w:tc')):
                p = tc.find(qn('w:p'))
                if p is None:
                    p = OxmlElement('w:p')
                    tc.insert(0, p)
                runs = p.findall(qn('w:r'))
                if not runs:
                    r = OxmlElement('w:r')
                    rPr = OxmlElement('w:rPr')
                    rf = OxmlElement('w:rFonts')
                    for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
                        rf.set(qn(a), CJK)
                    rPr.append(rf)
                    sz = OxmlElement('w:sz')
                    sz.set(qn('w:val'), '21')
                    rPr.append(sz)
                    r.append(rPr)
                    r.append(OxmlElement('w:t'))
                    p.append(r)
                else:
                    for r in runs:
                        if r.find(qn('w:rPr')) is None:
                            rPr = OxmlElement('w:rPr')
                            rf = OxmlElement('w:rFonts')
                            for a in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
                                rf.set(qn(a), CJK)
                            rPr.append(rf)
                            sz = OxmlElement('w:sz')
                            sz.set(qn('w:val'), '21')
                            rPr.append(sz)
                            r.insert(0, rPr)


# ---------- 弹窗提示 ----------
def gui_notify(title, message, is_error=False):
    try:
        import tkinter as tk
        from tkinter import messagebox
        root = tk.Tk()
        root.withdraw()
        if is_error:
            messagebox.showerror(title, message)
        else:
            messagebox.showinfo(title, message)
        root.destroy()
    except Exception:
        sys.stderr.write(f"[{title}] {message}\n")


# ---------- GUI ----------
def launch_gui():
    import tkinter as tk
    from tkinter import filedialog, messagebox, ttk

    root = tk.Tk()
    root.title("研究日志 MD → Word 生成器")
    root.geometry("640x200")
    root.resizable(False, False)

    md_var = tk.StringVar(value=DEFAULT_MD)
    out_var = tk.StringVar()

    def update_out():
        md = md_var.get().strip()
        if md and os.path.isfile(md):
            base, _ = os.path.splitext(md)
            out_var.set(base + '.docx')
        else:
            out_var.set('（请先选择有效的 .md 文件）')

    def choose_md():
        path = filedialog.askopenfilename(
            title="选择研究日志 Markdown 文件",
            filetypes=[("Markdown 文件", "*.md"), ("所有文件", "*.*")]
        )
        if path:
            md_var.set(path)
            update_out()

    def do_generate():
        md = md_var.get().strip()
        out = out_var.get().strip()
        if not md or not os.path.isfile(md):
            messagebox.showerror("无法生成", "请先选择有效的 .md 源文件。")
            return
        if not out or out.startswith('（'):
            base, _ = os.path.splitext(md)
            out = base + '.docx'
        try:
            parse_and_build(md, out)
            status_var.set(f"✅ 已生成：\n{out}")
            # 生成成功后约 1.5 秒自动关闭窗口（无需点确定）；出错才保留窗口
            root.after(1500, lambda: root.destroy() if root.winfo_exists() else None)
        except Exception:
            err = traceback.format_exc()
            status_var.set("❌ 生成失败，详见弹窗。")
            messagebox.showerror("生成失败", f"生成过程中出错：\n\n{err}")

    frm = ttk.Frame(root, padding=12)
    frm.pack(fill="both", expand=True)

    ttk.Label(frm, text="原始 Markdown 文件：").grid(row=0, column=0, sticky="w", pady=(0, 6))
    entry = ttk.Entry(frm, textvariable=md_var, width=60)
    entry.grid(row=1, column=0, columnspan=2, sticky="we", pady=(0, 4))
    ttk.Button(frm, text="浏览…", command=choose_md, width=10).grid(row=1, column=2, padx=(6, 0))

    ttk.Label(frm, text="输出 Word（与 md 同路径同名）：").grid(row=2, column=0, sticky="w", pady=(6, 2))
    ttk.Entry(frm, textvariable=out_var, width=60, state="readonly").grid(row=3, column=0, columnspan=2, sticky="we")
    ttk.Button(frm, text="生成", command=do_generate, width=10).grid(row=3, column=2, padx=(6, 0))

    status_var = tk.StringVar(value="选择 .md 文件后点击「生成」。")
    ttk.Label(frm, textvariable=status_var, foreground="#333", wraplength=560).grid(
        row=4, column=0, columnspan=3, sticky="w", pady=(10, 0))

    update_out()
    root.mainloop()


# ---------- 命令行 ----------
def headless_main():
    ap = argparse.ArgumentParser(description="研究日志 Markdown → Word 转换")
    ap.add_argument('md', nargs='?', default=DEFAULT_MD, help='输入 Markdown 路径')
    ap.add_argument('-o', '--output', default=None, help='输出 docx 路径（默认同目录同名）')
    args = ap.parse_args()

    md_path = os.path.abspath(args.md)
    if not os.path.exists(md_path):
        sys.stderr.write(f"[转换失败] 找不到输入文件：{md_path}\n")
        sys.exit(1)
    if args.output:
        docx_path = os.path.abspath(args.output)
    else:
        base, _ = os.path.splitext(md_path)
        docx_path = base + '.docx'

    try:
        out = parse_and_build(md_path, docx_path)
        sys.stderr.write(f"[转换完成] 已生成 Word 文档：{out}\n")
    except Exception:
        sys.stderr.write(f"[转换失败] 生成过程中出错：\n{traceback.format_exc()}\n")
        sys.exit(1)


if __name__ == '__main__':
    if len(sys.argv) > 1:
        headless_main()
    else:
        launch_gui()
